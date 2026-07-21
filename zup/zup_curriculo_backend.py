#!/usr/bin/env python3
"""Gera zup_curriculo_backend.docx para Candidato Nome Completo.

Variante especifica para vaga "Pessoa Desenvolvedora Backend Java" da Zup
(Greenhouse/LinkedIn, job-boards.greenhouse.io/zupinnovation), com foco em
modernizacao de sistema de comissionamento SAP. SAP e o sistema legado a
modernizar, nao a empresa contratante. A empresa e a Zup Innovation.

Diferencas vs gerar_curriculo_ats.py (generico):
- Perfil reescrito com keywords da vaga (Java, Spring, REST, integracao).
- Habilidades reordenadas com estrelas sendo backend JVM, REST, SQL, Camel,
  arquitetura, seguranca de APIs e CI/CD.
- Experiencia destaca modernizacao (Clean Architecture, refatoracao
  multi-tenant) e integracao robusta entre sistemas (gateway CNDs, GraphQL
  cliente, SOAP, RAG externo).
- Frontend recolhido a uma linha. Mobile e Go removidos.
- Secao IA removida (vaga nao pede). Cursos priorizam Spring, arquitetura e
  OAuth 2.0 (seguranca de APIs).

ATS-strict: uma coluna, sem tabelas, sem cores, sem foto. Fonte 10.5pt,
margens 2cm, line spacing 1.1 (alinhado a Alura). Sem em-dash/en-dash.

Conteudo canonico vem de briefing_llm_externo.md. Nada inventado.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Cm, Mm, RGBColor

OUTPUT = Path(__file__).parent / "zup_curriculo_backend.docx"

FONTE = "Calibri"
FONTE_NOME = 16
FONTE_H1 = 12
FONTE_H2 = 11
FONTE_CORPO = 10

# Flag de cor: True = layout verde de uma coluna (estilo Greenhouse-friendly).
#              False = volta ao ATS estrito preto e branco.
# Greenhouse (parser da Zup) le texto colorido sem problema. Se um portal
# legado barrar, basta mudar para False e regenerar.
USE_COR = True

# Verde floresta ~#2E8B57. Contrastante, profissional, ATS-safe em parsers
# modernos. RGB direct (46, 139, 87).
COR_DESTAQUE = RGBColor(0x2E, 0x8B, 0x57)
COR_CORPO = RGBColor(0x00, 0x00, 0x00)


def _style(doc: Document) -> None:
    """A4, margens 1.2cm (compromisso ATS) e fonte padrao."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(FONTE_CORPO)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.05


def _cor_ambiente() -> RGBColor | None:
    """Retorna a cor de destaque se USE_COR ligado, senao None (preto default)."""
    return COR_DESTAQUE if USE_COR else None


def _nome(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_NOME)
    r.font.name = FONTE
    cor = _cor_ambiente()
    if cor is not None:
        r.font.color.rgb = cor


def _cargo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_H1)
    r.font.name = FONTE
    cor = _cor_ambiente()
    if cor is not None:
        r.font.color.rgb = cor


def _contato(doc: Document, partes: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(" | ".join(partes))
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _add_borda_inferior(paragraph, cor_hex: str = "2E8B57", tamanho: str = "6") -> None:
    """Adiciona borda inferior ao paragrafo via XML direto (python-docx nao tem API)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), tamanho)  # 1/8 pt. 6 = 0.75pt.
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), cor_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _h2(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_H2)
    r.font.name = FONTE
    if USE_COR:
        cor = _cor_ambiente()
        if cor is not None:
            r.font.color.rgb = cor
        _add_borda_inferior(p)


def _paragrafo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _bullet(doc: Document, texto: str, negrito_prefixo: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    if negrito_prefixo:
        r1 = p.add_run(negrito_prefixo)
        r1.bold = True
        r1.font.size = Pt(FONTE_CORPO)
        r1.font.name = FONTE
    r2 = p.add_run(texto)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def _linha_data(doc: Document, esquerda: str, direita: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(18.6), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(esquerda)
    r1.bold = True
    r1.font.size = Pt(FONTE_CORPO)
    r1.font.name = FONTE
    r2 = p.add_run("\t" + direita)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def construir() -> Document:
    doc = Document()
    _style(doc)

    # ---- Cabecalho ----
    _nome(doc, "Candidato Nome Completo")
    _cargo(doc, "Desenvolvedor Backend Java")
    _contato(
        doc,
        [
            "(00) 00000-0000",
            "candidato.exemplo@dominio.com",
            "Cidade, UF",
            "linkedin.com/in/seu-perfil/",
            "github.com/seu-usuario",
            "seu-portfolio.com/portfolio/seu-usuario",
        ],
    )

    # ---- Perfil (focado nas keywords da vaga) ----
    _h2(doc, "Perfil")
    _paragrafo(
        doc,
        "Desenvolvedor Backend Java com 5 anos de experiência em Spring Boot e APIs "
        "RESTful (Representational State Transfer), com passagem por sete domínios "
        "de negócio distintos (agronegócio, fintech, logística, automotivo, "
        "gamificação corporativa, compliance e educação). Foco em padrões de "
        "arquitetura backend (Clean Architecture, CQRS, DDD), integração robusta "
        "entre sistemas via Apache Camel e filas, otimização de bancos de dados "
        "sobre PostgreSQL e MySQL, e monitoramento de aplicações com "
        "observabilidade. Bacharel em Ciência da Computação pela UFPA (2024).",
    )

    # ---- Habilidades (reordenadas para a vaga) ----
    _h2(doc, "Habilidades")
    _bullet(doc, "Java (11/16/17/21), Spring Boot (2.x e 3.x), Spring Data JPA, Spring Security, Spring Cloud (Eureka, Feign), Spring WebFlux.", "Backend (JVM): ")
    _bullet(doc, "REST (Representational State Transfer), OpenAPI/Swagger, SOAP/WSDL (integração legada), GraphQL (cliente via graphql-java-generator), WebClient.", "APIs e contratos: ")
    _bullet(doc, "PostgreSQL, MySQL, Hibernate/JPA, Hibernate Envers (auditoria), hibernate-spatial + JTS (geometria), Flyway (migrations), QueryDSL, jOOQ (Live), SQL nativo para otimização de performance.", "Bancos relacionais: ")
    _bullet(doc, "Apache Camel (roteamento e integração B2B), AWS SQS, Redis (cache), Valkey, WebSocket (STOMP/SockJS).", "Integração e mensageria: ")
    _bullet(doc, "Clean Architecture, CQRS-lite, DDD (estratégico e tático), padrões de projeto (Strategy, Chain of Responsibility, Template Method, State, Command), SOLID, modular architecture.", "Arquitetura e padrões: ")
    _bullet(doc, "JWT (jjwt), Keycloak (OAuth2/OIDC, OpenID Connect), AWS Cognito. OAuth 2.0 estudado a fundo (Casa do Código).", "Segurança de APIs: ")
    _bullet(doc, "Git, GitLab CI/CD, Jenkins, Docker, Portainer, Terraform. Pipelines de CI/CD com automação de deploys. AWS (ECS, ECR, S3, CloudFront, Lambda, Elastic Beanstalk, RDS).", "DevOps e CI/CD: ")
    _bullet(doc, "JUnit, Mockito, AssertJ, Testcontainers.", "Testes: ")
    _bullet(doc, "Sentry, Logstash, OpenTelemetry para monitoramento de aplicações em produção. Kanban (Redmine, ClickUp).", "Observabilidade e metodologia: ")
    _bullet(doc, "Angular, React e Vue para atuação full stack quando necessário.", "Frontend (complementar): ")

    # ---- Soft Skills (Alura recomenda 4-6) ----
    _h2(doc, "Soft Skills e Idiomas")
    _bullet(doc, "Comunicação técnica estruturada a pares: coleta de requisitos com cliente, tradução de problemas de negócio em contratos de API e UML de domínio.")
    _bullet(doc, "Pensamento crítico na revisão técnica de código e na modelagem de domínios complexos (inspeções rodoviárias, exames de saúde, apontamento de horas).")
    _bullet(doc, "Colaboração em cultura de testes e revisão por pares (Testcontainers, AssertJ, Mockito).")
    _bullet(doc, "Português nativo. Inglês: leitura técnica fluente, curso em andamento (2025).", "Idiomas: ")

    # ---- Experiencia ----
    _h2(doc, "Experiência")

    # iUsecase
    _linha_data(doc, "iUsecase Tecnologia e Inovação", "Jul 2025 - Atual")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Desenvolvedor Backend Pleno 1 (cargo CTPS). Atuação remota em três produtos com foco em design de código, arquitetura limpa e integração de sistemas.")
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(doc, "Liderei modernização de sistema legado de fiscalização rodoviária migrando para Clean Architecture com CQRS-lite (Commands via ports, leitura por QueryServices @Cacheable), Java 21 e Spring Boot 3.2. Otimizei queries críticas com SQL nativo PostgreSQL (window functions, full-text com unaccent PT-BR, UPSERT atômico) e isolei sequenciais concorridos com locks pessimistas (PESSIMISTIC_WRITE). Adicionei testes de segurança IDOR cobrindo controllers e teste de concorrência no sequencial, com cobertura ampla via Testcontainers, AssertJ e Mockito.", "Consol: ")
    _bullet(doc, "Modelei arquitetura multi-tenant database-driven para timesheet com policies configuráveis em banco e resolvers @Cacheable em Caffeine, em arquitetura hexagonal documentada em C4. Liderei migração para eliminar condicionais de cliente do core: módulo de timelogging migrado, demais módulos ainda pendentes (dívida documentada). Spring Boot 3.2 com Java 17 e QueryDSL para queries dinâmicas tipadas sobre PostgreSQL.", "Apontamento: ")
    _bullet(doc, "Integrei serviço externo de IA sobre exames em sistema de saúde em arquitetura hexagonal, com upload de PDF, jobs assíncronos (@Async, Quartz para refresh de tokens de parceiros) e eventos pós-commit (@TransactionalEventListener AFTER_COMMIT para WhatsApp e sync de links). Backend Spring Boot 3.2 com Java 21 e jOOQ para dashboards. O backend de IA é operado por terceiros; meu trabalho foi a orquestração e a integração robusta entre sistemas distintos.", "Live2U: ")
    _bullet(doc, "Mantive pipelines de CI/CD em GitLab com deploy de imagens ARM64 para ECS via ECR (backend) e S3 com CloudFront (frontend). Colaborei em cultura de testes com Testcontainers, AssertJ e Mockito, revisão técnica de código no time e suporte a aplicações em produção com diagnóstico de defeitos.")

    # itexto
    _linha_data(doc, "itexto Consultoria em Tecnologia", "Out 2021 - Abr 2025")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Programador (cargo CTPS, CBO 3171-10). Função Full Stack no ciclo completo de software em sete domínios de negócio (seis em Java e um em Go).")
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(doc, "Construí plataforma Ativus de crédito rural multi-tenant em Spring Boot 2.3 e Java 11, com separação Command/Query em agregados DDD e SQL nativo para stored procedures multi-schema em PostgreSQL. Orquestrei integrações assíncronas com @Async e cron jobs (Serasa, gateway Vindi, SendGrid) e gateway de CNDs governamentais via Apache Camel e AWS SQS. Implementei barter agrícola da Corteva em Spring Cloud (Eureka, Feign) e Keycloak OIDC. Cobertura ampla de testes com JUnit e Mockito.", "Agronegócio: ")
    _bullet(doc, "Implementei tokenizadora de ativos (QR-Capital) com pipeline transacional auditável baseado em Transaction Log/Outbox (unidades REQUIRES_NEW + REPEATABLE_READ com rastreabilidade via ThreadLocal). Cliente GraphQL via graphql-java-generator + WebClient, com Spring Data JPA, Hibernate Envers sobre PostgreSQL e OAuth2/Keycloak. Orquestração assíncrona em Apache Camel + SQS para sync de carteiras e status de transferências bancárias.", "Fintech: ")
    _bullet(doc, "Desenvolvi marketplace de frete Flex-Frete com cotação, contratação e notas fiscais. Usei TransactionTemplate programático para transações financeiras de carteira (escopo condicional débito/crédito) e cron job de expiração/cancelamento de fretes. Spring Boot 2.5, PostgreSQL, Redis (cache), React 17.", "Logística: ")
    _bullet(doc, "Estruturei monorepo Weex com microserviço async dedicado (weex.async) com rotas Apache Camel + SQS para certificados, expurgo, logs e processamento paralelo de imagens (CompletableFuture). Geração de certificados em AWS Lambda e IaC com Terraform. Cobertura ampla de testes de integração ponta-a-ponta com @SpringBootTest.", "Gamificação corporativa: ")
    _bullet(doc, "Mantive plataforma Wirelist (Starcom) de documentação técnica para montadoras com SQL nativo (subqueries de aprovação, window functions implícitas) e transações isoladas para importação em lotes (REQUIRES_NEW + REPEATABLE_READ). Spring Boot 2.2, MySQL, Elastic Beanstalk, Angular 15.", "Automotivo/industrial: ")
    _bullet(doc, "Participei do ciclo completo de sete sistemas em três anos e meio (incluindo RRZ de compliance com auditoria via Hibernate Envers): coleta de requisitos com o cliente, contratos de API em Swagger, UML das classes de domínio, desenvolvimento backend (Java/Spring), integrações com APIs externas e S3, Apache POI para planilhas, SQL nativo para otimização de performance quando JPQL não resolvia, testes com JUnit e Mockito, deploy com Docker e Portainer, suporte a aplicações em produção com diagnóstico de defeitos.", "Transversais: ")

    # ---- Formacao ----
    _h2(doc, "Formação")
    _linha_data(doc, "Bacharelado em Ciência da Computação: UFPA", "Abr 2017 - Jan 2024")
    _paragrafo(doc, "Concluído em janeiro de 2024. Diploma emitido em julho de 2024 (Belém/PA).")

    # ---- Formacao complementar (prioriza Spring, arquitetura e seguranca) ----
    _h2(doc, "Formação Complementar")
    _bullet(doc, "Udemy (out/2021). Microservices do 0 com Spring Cloud, Spring Boot e Docker (Feign, Eureka, Circuit Breaker, Resilience4j).", "Spring/Cloud: ")
    _bullet(doc, "Casa do Código. OAuth 2.0: fluxos de autorização, tokens, escopos e boas práticas de segurança de APIs.", "Segurança de APIs: ")
    _bullet(doc, "Dev Eficiente (2025). Jornada Dev Eficiente: DDD, system design, escalabilidade, CDD, resiliência.", "Arquitetura: ")
    _bullet(doc, "Joshua Bloch. Estudo com anotações em Notion por capítulo.", "Effective Java: ")

    return doc


# ---- Assertions TDD inline ----
def validar(doc: Document) -> None:
    """Valida restricoes ATS e cobertura de keywords da vaga."""
    texto = "\n".join(p.text for p in doc.paragraphs)

    # 1. Estrutura: secoes obrigatorias presentes.
    secoes = ["PERFIL", "HABILIDADES", "SOFT SKILLS E IDIOMAS", "EXPERIÊNCIA", "FORMAÇÃO", "FORMAÇÃO COMPLEMENTAR"]
    for s in secoes:
        assert s in texto, f"Secao obrigatoria ausente: {s}"

    # 1b. Idiomas deve aparecer (na seção Soft Skills e Idiomas).
    assert "Idiomas:" in texto, "Bullet de Idiomas ausente"

    # 2. Sem tabelas de layout (ATS nao parseia bem).
    assert len(doc.tables) == 0, f"ATS proibe tabelas de layout. Encontradas: {len(doc.tables)}"

    # 3. Keywords da vaga (obrigatorias) para match ATS.
    keywords_vaga = [
        "Java", "Spring", "REST", "PostgreSQL", "MySQL",
        "Clean Architecture", "Apache Camel", "Git", "CI/CD",
        "Docker", "OAuth", "API",
        # Literais da JD que faltavam:
        "pipelines", "monitoramento",
        # Padrões técnicos reais (defensáveis em entrevista):
        "SQL nativo", "REQUIRES_NEW", "Testcontainers",
    ]
    for kw in keywords_vaga:
        assert kw in texto, f"Keyword da vaga ausente: {kw}"

    # 3d. Anti-precisao-falsa: proibe contagens inventadas ou irrelevantes.
    # Numeros como "32 queries", "7 testes", "1h TTL" soam como falsa precisao.
    proibidos_precisao = [
        "32 queries", "16 repositórios", "7 testes", "TTL 1h",
        "123 commands", "111 queries", "13 métodos", "10 cron",
        "720 classes", "537", "629 classes", "14 rotas", "5 unidades",
        "3 locks",
    ]
    for termo in proibidos_precisao:
        assert termo not in texto, f"Métrica de falsa precisão (cortar): {termo}"

    # 3c. Anti-invencao: termos que NAO devem aparecer (sem evidencia em codigo).
    proibidos = [
        "Redis pub/sub",  # só cache, sem pub/sub real
        "WireMock",  # no pom mas sem uso
    ]
    for termo in proibidos:
        assert termo not in texto, f"Termo proibido (sem evidencia): {termo}"

    # 3b. Siglas expandidas ao menos uma vez (regra Alura).
    siglas_expandidas = [
        "Representational State Transfer",
        "OpenID Connect",
    ]
    for sigla in siglas_expandidas:
        assert sigla in texto, f"Sigla nao expandida: {sigla}"

    # 4. Contato presente.
    assert "(00) 00000-0000" in texto, "Telefone ausente"
    assert "candidato.exemplo@dominio.com" in texto, "E-mail ausente"

    # 5. Cargos CTPS explicitados (honestidade).
    assert "Desenvolvedor Backend Pleno 1" in texto, "Cargo CTPS iUsecase ausente"
    assert "CBO 3171-10" in texto, "Cargo CTPS itexto ausente"

    # 6. Regra de honestidade (regra 4 do handoff): jOOQ so no Live2U.
    assert "jOOQ" in texto, "jOOQ deve aparecer (Live2U)"
    for p in doc.paragraphs:
        if "Consol:" in p.text or "Apontamento:" in p.text:
            assert "jOOQ" not in p.text, f"jOOQ nao deve aparecer em {p.text[:40]}..."

    # 7. Sem em-dashes (regra de estilo do AGENTS.md).
    assert "—" not in texto, "Em-dash encontrado. Usar pontos, virgulas ou dois-pontos."
    assert "–" not in texto, "En-dash encontrado. Usar hifen simples com espacos."

    # 8. Verbo de acao no inicio de cada bullet de experiencia.
    verbos_aceitos = {
        "Desenvolvi", "Modelei", "Integrei", "Mantive", "Colaborei",
        "Construí", "Implementei", "Estruturei", "Participei",
        "Liderei", "Otimizei", "Adicionei", "Orquestrei", "Usei",
    }
    prefixos_experiencia = {
        "Consol:", "Apontamento:", "Live2U:",
        "Agronegócio:", "Fintech:", "Logística:", "Gamificação corporativa:",
        "Automotivo/industrial:", "Compliance:", "Transversais:",
    }
    for p in doc.paragraphs:
        for prefixo in prefixos_experiencia:
            if p.text.startswith(prefixo):
                corpo = p.text[len(prefixo):].strip()
                palavras = corpo.replace(".", " ").split()
                tem_verbo = any(w.rstrip(",.;:") in verbos_aceitos for w in palavras[:3])
                assert tem_verbo, (
                    f"Bullet '{prefixo}' deve comecar com verbo de acao no passado "
                    f"(entre os 3 primeiros tokens). Inicio: '{palavras[0] if palavras else ''}'."
                )

    # 9. Anti-invencao: termos que NAO devem aparecer (nao tenho SAP/K8s reais).
    #    Permite mencao honesta em secao de desejaveis, mas proibe afirmar como
    #    experiencia direta. Aqui so bloqueia afirmacao explicita.
    for termo_proibido in ["experiência em SAP", "SAP ABAP", "Kubernetes em produção"]:
        assert termo_proibido not in texto, f"Termo proibido (invencao): {termo_proibido}"

    print("[OK] Validacao ATS + cobertura de keywords da vaga passou.")
    print(f"  - Paragrafos: {len(doc.paragraphs)}")
    print(f"  - Tabelas: {len(doc.tables)} (deve ser 0)")
    print(f"  - Caracteres: {len(texto)}")


def main() -> None:
    doc = construir()
    validar(doc)
    doc.save(str(OUTPUT))
    print(f"[OK] Salvo em: {OUTPUT}")


if __name__ == "__main__":
    main()
