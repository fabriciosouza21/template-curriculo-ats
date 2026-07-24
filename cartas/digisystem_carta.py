#!/usr/bin/env python3
"""Gera carta de apresentacao para a vaga fullstack Java/Angular da Digisystem.

Carta narrativa (nao currículo). Texto corrido de 5 paragrafos dirigido à
equipe de recrutamento, com cabecalho de contato. Conteudo fiel ao
briefing e aos cases canonicos dos YAMLs em data/. Nada inventado.

ATS-safe: uma coluna, sem tabelas, sem foto. A4, fonte 10.5pt corpo,
margens 1.2cm. Sem em-dashes (U+2014) nem en-dashes (U+2013).

Uso:
    python3 cartas/digisystem_carta.py
Saida: cartas/digisystem_carta.docx e conversao para PDF.
"""
from __future__ import annotations

import re
import shutil
import subprocess
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Mm, Pt, RGBColor

OUTPUT = Path(__file__).parent / "digisystem_carta.docx"

FONTE = "Calibri"
FONTE_NOME = 16
FONTE_CORPO = 10.5
FONTE_DATA = 10

# Mesmo azul profundo do manifesto Digisystem (#0B1641).
COR_DESTAQUE = RGBColor(0x0B, 0x16, 0x41)


def _style(doc: Document) -> None:
    """A4, margens 1.2cm e fonte padrao."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(FONTE_CORPO)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.05


def _nome(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_NOME)
    r.font.name = FONTE
    r.font.color.rgb = COR_DESTAQUE


def _contato(doc: Document, partes: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(" | ".join(partes))
    r.font.size = Pt(FONTE_DATA)
    r.font.name = FONTE


def _paragrafo(doc: Document, texto: str, espaco_depois: int = 6,
               justificado: bool = True) -> None:
    p = doc.add_paragraph()
    if justificado:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(espaco_depois)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _destinatario(doc: Document, linhas: list[str]) -> None:
    for linha in linhas:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        r = p.add_run(linha)
        r.font.size = Pt(FONTE_DATA)
        r.font.name = FONTE


def construir() -> Document:
    doc = Document()
    _style(doc)

    # ---- Cabecalho ----
    _nome(doc, "Candidato Nome Completo")
    _contato(
        doc,
        [
            "(00) 00000-0000",
            "candidato.exemplo@dominio.com",
            "Cidade, UF",
            "linkedin.com/in/seu-perfil/",
            "github.com/seu-usuario",
        ],
    )

    # ---- Destinatario e data ----
    _destinatario(
        doc,
        [
            "À equipe de recrutamento da Digisystem",
            "Cidade, UF, 22 de julho de 2026.",
        ],
    )

    # ---- Saudacao ----
    _paragrafo(doc, "Prezados,", espaco_depois=6, justificado=False)

    # ---- Paragrafo 1: abertura + por que a vaga ----
    _paragrafo(
        doc,
        "Escrevo para manifestar interesse na vaga de Desenvolvedor Full "
        "Stack com Java e Angular na Digisystem. Sou Desenvolvedor Full "
        "Stack Pleno com quase cinco anos de experiência em entregas "
        "end-to-end em Java/Spring no backend e Angular e React no "
        "frontend, com dez sistemas em produção em sete domínios de "
        "negócio distintos. A combinação de Spring Boot com JPA/Hibernate "
        "no backend, Angular com TypeScript no frontend e APIs REST é "
        "exatamente a pilha que venho operando no dia a dia, o que torna a "
        "vaga um encaixe natural para o que já entrego hoje.",
    )

    # ---- Paragrafo 2: cases de impacto (backend + frontend) ----
    _paragrafo(
        doc,
        "Nos últimos anos concentrei trabalho em sistemas com volume e "
        "regras de negócio relevantes. Na itexto construí por três anos e "
        "meio a plataforma Weex de gamificação corporativa, que processou "
        "mais de 2,4 milhões de ações e cerca de 290 mil inscrições em "
        "eventos, em arquitetura de microsserviços com mensageria "
        "assíncrona. Na mesma empresa desenvolvi features de crédito rural "
        "na plataforma multi-tenant Ativus, que atende 79 clientes do "
        "agronegócio, incluindo o gateway de Certidões Negativas de "
        "Débitos que acelerou a validação cadastral integrando Receita "
        "Federal, estaduais e Ibama. Na iUsecase, onde atuo atualmente, "
        "desenvolvi backend para fiscalização de malha rodoviária em Minas "
        "Gerais com sincronização mobile offline-first sobre 14 rodovias, "
        "usando lock distribuído para eliminar perda de dados em campo.",
    )

    # ---- Paragrafo 3: diferenciais da vaga ----
    _paragrafo(
        doc,
        "Sobre os diferenciais mencionados no anúncio, tenho trajetória "
        "direta com a maioria deles. Trabalho com Docker na rotina de "
        "desenvolvimento e deploy, CI/CD em GitLab com publicação de "
        "imagens para ECS via ECR e Jenkins, e nuvem AWS com serviços como "
        "S3, SQS, SES, ECS, ECR, CloudFront, Cognito, Lambda e RDS. "
        "Construo APIs REST com contratos em OpenAPI e Swagger, e mantenho "
        "cultura de testes automatizados com JUnit, Mockito, AssertJ e "
        "Testcontainers no backend, e Cypress e Karma/Jasmine no frontend "
        "Angular. Microserviços são parte recorrente do meu trabalho há "
        "anos, com mensageria em Apache Camel e AWS SQS.",
    )

    # ---- Paragrafo 4: bancos + metodologia + honestidade ----
    _paragrafo(
        doc,
        "Em bancos relacionais tenho domínio de PostgreSQL e MySQL com "
        "Hibernate/JPA, Flyway para migrations, QueryDSL para queries "
        "dinâmicas e SQL nativo quando preciso de desempenho. Trabalho em "
        "ambientes ágeis com Kanban, faço code review com padrões "
        "definidos pelo time e análise de causa raiz de incidentes em "
        "produção. Para ser transparente: não tenho experiência "
        "consolidada com Kubernetes, que aparece como diferencial no "
        "anúncio, embora tenha familiaridade com os conceitos de "
        "orquestração de contêineres. Prefiro deixar isso claro a "
        "prometer o que não entregaria.",
    )

    # ---- Paragrafo 5: fechamento ----
    _paragrafo(
        doc,
        "Sou bacharel em Ciência da Computação pela UFPA (2024) e atuo "
        "remoto desde 2021, o que me dá maturidade em comunicação "
        "assíncrona e autonomia. Disponho-me a uma conversa para detalhar "
        "como posso contribuir com os times da Digisystem, e agradeço "
        "desde já a atenção dedicada a esta candidatura.",
        espaco_depois=10,
    )

    # ---- Assinatura ----
    _paragrafo(doc, "Atenciosamente,", espaco_depois=0, justificado=False)
    _paragrafo(
        doc, "Candidato Nome Completo", espaco_depois=0, justificado=False
    )

    return doc


def validar(doc: Document) -> None:
    """Valida restricoes da carta: sem em-dashes, sem tabelas, keywords presentes."""
    texto = "\n".join(p.text for p in doc.paragraphs)

    assert len(doc.tables) == 0, f"ATS proibe tabelas. Encontradas: {len(doc.tables)}"

    assert "—" not in texto, "Em-dash (U+2014) encontrado."
    assert "–" not in texto, "En-dash (U+2013) encontrado."

    keywords = [
        "Java", "Spring", "Angular", "TypeScript", "REST",
        "PostgreSQL", "MySQL", "JPA", "Hibernate", "Docker",
        "CI/CD", "AWS", "microsserviços", "JUnit", "Mockito",
        "Cypress", "GitLab", "OpenAPI",
    ]
    faltando = [kw for kw in keywords if kw not in texto]
    assert not faltando, f"Keywords da vaga ausentes: {faltando}"

    # Metricas reais do briefing (nao inventar precisao adicional).
    assert "2,4 milhões" in texto, "Métrica Weex ausente."
    assert "290 mil" in texto, "Métrica Weex inscrições ausente."
    assert "79 clientes" in texto, "Métrica Ativus ausente."
    assert "14 rodovias" in texto, "Métrica Consol ausente."

    # Honestidade: Kubernetes declarado como não consolidado.
    assert "não tenho experiência consolidada com Kubernetes" in texto, (
        "Carta deve declarar honestamente ausência de Kubernetes."
    )

    assert "À equipe de recrutamento da Digisystem" in texto
    assert "(00) 00000-0000" in texto
    assert "candidato.exemplo@dominio.com" in texto

    # Carta precisa caber em 1 pagina.
    n_paragrafos = len(doc.paragraphs)
    assert 12 <= n_paragrafos <= 30, (
        f"Carta com {n_paragrafos} parágrafos fora do esperado (12-30)."
    )

    print("[OK] validação da carta passou.")
    print(f"  - Parágrafos: {n_paragrafos}")
    print(f"  - Tabelas: {len(doc.tables)} (deve ser 0)")
    print(f"  - Caracteres: {len(texto)}")


def _converter_pdf(docx_path: Path) -> Path:
    if shutil.which("soffice") is None and shutil.which("libreoffice") is None:
        raise RuntimeError("soffice/libreoffice não encontrado no PATH.")
    soffice = shutil.which("soffice") or shutil.which("libreoffice")
    subprocess.run(
        [soffice, "--headless", "--convert-to", "pdf",
         "--outdir", str(docx_path.parent), str(docx_path)],
        capture_output=True, text=True, timeout=120,
    )
    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.exists():
        raise RuntimeError("conversão para PDF falhou.")
    return pdf_path


def _contar_paginas_pdf(pdf_path: Path) -> int:
    with open(pdf_path, "rb") as f:
        data = f.read()
    counts = re.findall(rb"/Count\s+(\d+)", data)
    if counts:
        return max(int(c) for c in counts)
    return len(re.findall(rb"/Type\s*/Page[^s]", data))


def main() -> int:
    print("[1/4] Montando carta...")
    doc = construir()

    print("[2/4] Validando regras...")
    validar(doc)

    print(f"[3/4] Salvando em {OUTPUT}...")
    doc.save(str(OUTPUT))

    print("[4/4] Convertendo para PDF...")
    try:
        pdf_path = _converter_pdf(OUTPUT)
    except RuntimeError as exc:
        print(f"[WARN] {exc}", file=sys.stderr)
        print(f"[OK] gerado: {OUTPUT}")
        return 0

    n_paginas = _contar_paginas_pdf(pdf_path)
    print(f"  - Páginas: {n_paginas}")
    if n_paginas > 1:
        print(
            f"[FAIL] carta tem {n_paginas} páginas, excede o limite de 1.",
            file=sys.stderr,
        )
        return 1

    print(f"[OK] gerado: {OUTPUT}")
    print(f"  - PDF: {pdf_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
