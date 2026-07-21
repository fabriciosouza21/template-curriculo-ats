#!/usr/bin/env python3
"""Gera curriculo_ats.docx para Candidato Nome Completo.

Layout ATS estrito: uma coluna, sem tabelas, sem cores, sem foto, texto puro.
Baseado em briefing_llm_externo.md (conteudo canonico) e handoff (regras de
honestidade). Diretrizes visuais seguem artigos da Alura ("Como fazer
curriculo" e "IA para fazer curriculo"): fonte 10-12pt, margens ~2cm, verbos
de acao nos bullets, sem em-dashes, secao de idiomas.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Mm

OUTPUT = Path(__file__).parent / "curriculo_ats.docx"

# Fonte ATS-safe: Calibri fallback universal. Tamanhos alinhados a Alura
# (10-12pt) com margens ~2cm. Aceita 3 paginas para perfis com bagagem.
FONTE = "Calibri"
FONTE_NOME = 18
FONTE_H1 = 12
FONTE_H2 = 11
FONTE_CORPO = 10.5


def _style(doc: Document) -> None:
    """Configura A4, margens ~2cm (recomendacao Alura) e fonte padrao."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(FONTE_CORPO)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.1


def _nome(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_NOME)
    r.font.name = FONTE


def _cargo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_H1)
    r.font.name = FONTE


def _contato(doc: Document, partes: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(" | ".join(partes))
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _h2(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_H2)
    r.font.name = FONTE
    # Borda inferior discreta via underline em run separado seria ruido para
    # ATS. Mantemos apenas bold + caixa alta.


def _paragrafo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _bullet(doc: Document, texto: str, negrito_prefixo: str = "") -> None:
    """Bullet com prefixo em negrito opcional. ATS lê bullet char '-' bem."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    if negrito_prefixo:
        r1 = p.add_run(negrito_prefixo)
        r1.bold = True
        r1.font.size = Pt(FONTE_CORPO)
        r1.font.name = FONTE
    r2 = p.add_run(texto)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def _linha_data(doc: Document, esquerda: str, direita: str) -> None:
    """Linha com cargo a esquerda e data a direita via tab stop."""
    from docx.enum.text import WD_TAB_ALIGNMENT

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(esquerda)
    r1.bold = True
    r1.font.size = Pt(FONTE_CORPO)
    r1.font.name = FONTE
    r2 = p.add_run("\t" + direita)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def construir() -> Document:
    doc = Document()
    _style(doc)

    # ---- Cabecalho ----
    _nome(doc, "Candidato Nome Completo")
    _cargo(doc, "Engenheiro de Software Pleno")
    _contato(
        doc,
        [
            "(00) 00000-0000",
            "candidato.exemplo@dominio.com",
            "Cidade, UF",
            "linkedin.com/in/seu-perfil/",
            "github.com/seu-usuario",
            "seu-portfolio.com/portfolio/seu-usuario",
        ],
    )

    # ---- Perfil (estilo Alura: enxuto, orientado a conquista) ----
    _h2(doc, "Perfil")
    _paragrafo(
        doc,
        "Engenheiro de Software Pleno com 5 anos de experiência em desenvolvimento Full Stack "
        "focado em Java/Spring, Clean Architecture e design de código. Passagem por sete domínios "
        "de negócio distintos (agronegócio, fintech, logística, automotivo, gamificação corporativa, "
        "compliance e educação). Bacharel em Ciência da Computação pela UFPA (2024).",
    )

    # ---- Habilidades ----
    _h2(doc, "Habilidades")
    _bullet(doc, "Java (11/16/17/21), Spring Boot (2.x e 3.x), Spring Data JPA, Spring Security, Spring Cloud (Eureka, Feign), Spring WebFlux.", "Backend (JVM): ")
    _bullet(doc, "PostgreSQL, MySQL, Hibernate/JPA, Hibernate Envers, hibernate-spatial + JTS, Flyway, QueryDSL, jOOQ (Live), SQL nativo.", "Persistência: ")
    _bullet(doc, "Apache Camel, AWS SQS, Redis (cache e pub/sub), Valkey, WebSocket (STOMP/SockJS).", "Mensageria: ")
    _bullet(doc, "Angular (7 a 20), React (17/18/19), Vue 2, React Native/Expo, Flutter, Vite, MapLibre/Leaflet.", "Frontend: ")
    _bullet(doc, "JWT (jjwt), Keycloak (OAuth2/OIDC), AWS Cognito.", "Auth: ")
    _bullet(doc, "AWS (S3, SQS, SES, ECS, ECR, CloudFront, Cognito, Lambda, Elastic Beanstalk, RDS), Cloudinary, Docker, Portainer, Terraform, GitLab CI/CD, Jenkins.", "Cloud/DevOps: ")
    _bullet(doc, "JUnit, Mockito, AssertJ, Testcontainers, WireMock, Cypress, Karma/Jasmine, Vitest, Jest.", "Testes: ")
    _bullet(doc, "REST, OpenAPI/Swagger, GraphQL (cliente via graphql-java-generator), SOAP/WSDL.", "APIs e contratos: ")
    _bullet(doc, "Apache POI (Excel), iText/PDFBox (PDF), OpenCSV, Velocity.", "Documentos/dados: ")
    _bullet(doc, "Sentry, Logstash, OpenTelemetry. Kanban (Redmine, ClickUp).", "Observabilidade e metodologia: ")

    # ---- Experiencia ----
    _h2(doc, "Experiência")

    # iUsecase
    _linha_data(doc, "iUsecase Tecnologia e Inovação", "Jul 2025 - Atual")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Desenvolvedor Backend Pleno 1 (cargo CTPS). Atuação remota em três produtos com foco em design de código e arquitetura limpa.")
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(doc, "case principal. Desenvolvi backend em Spring Boot 3.2 e Java 21 para fiscalização de malha rodoviária em Clean Architecture com CQRS-lite (leitura por QueryService com @Cacheable, escrita por Commands via ports), usando Spring Data JPA com hibernate-spatial/JTS para geometria rodoviária. Implementei sincronização offline-first mobile com lock distribuído em Redis e fallback degradado (consol-sync-api em Node/Express + WebSocket + Valkey), frontend em React 19 com Vite e mapas MapLibre, auth em Cognito e mobile em Flutter.", "Consol: ")
    _bullet(doc, "Modelei arquitetura multi-tenant database-driven para timesheet com policies configuráveis em banco (policy, policy_rule, tenant_policy) e resolvers @Cacheable em Caffeine TTL 1h. Migrei o módulo de timelogging para eliminar condicionais de cliente; Initiative/Workbook/Sprint ainda pendentes (dívida documentada). Spring Boot 3.2 com Java 17, QueryDSL e Angular 17.", "Apontamento: ")
    _bullet(doc, "Integrei serviço externo de RAG sobre exames no frontend Angular, com upload de PDF, jobs assíncronos com polling e respostas com citações atreladas às observações de cada exame. Backend Spring Boot 3.2 com Java 21, jOOQ e integrações com serviços de saúde. O backend de IA é externo (operado pela Sys3); meu trabalho foi a orquestração e integração.", "Live2U: ")
    _bullet(doc, "Mantive CI/CD em GitLab com deploy de imagens ARM64 para ECS via ECR (backend) e S3 com CloudFront (frontend). Colaborei em cultura de testes com Testcontainers, AssertJ e WireMock, revisão de código no time e suporte a aplicações em produção.")

    # itexto
    _linha_data(doc, "itexto Consultoria em Tecnologia", "Out 2021 - Abr 2025")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Programador (cargo CTPS, CBO 3171-10). Função Full Stack no ciclo completo de software em sete domínios de negócio (seis em Java e um em Go).")
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(doc, "Construí plataforma Ativus de crédito rural e análise de crédito em Spring Boot 2.3, Java 11, MySQL, Flyway e Angular 13, com gateway de CNDs governamentais em Apache Camel e SQS. Implementei sistema de barter agrícola da Corteva em Spring Cloud (Eureka, Feign), Vue 2, PostgreSQL com Hibernate Envers e Keycloak OIDC.", "Agronegócio: ")
    _bullet(doc, "Implementei tokenizadora de ativos (QR-Capital) como cliente GraphQL de API externa (graphql-java-generator + WebClient), com Spring Data JPA e Hibernate Envers sobre PostgreSQL, OAuth2/Keycloak, Spring Boot 2.7/Java 16.", "Fintech: ")
    _bullet(doc, "Desenvolvi marketplace de frete Flex-Frete com cotação, contratação, notas fiscais e mensageria em Camel e SQS. Spring Boot 2.5, PostgreSQL, Redis, React 17.", "Logística: ")
    _bullet(doc, "Estruturei monorepo Weex (bwell) com múltiplos microsserviços, processamento assíncrono em Camel e SQS, geração de certificados em AWS Lambda e IaC com Terraform.", "Gamificação corporativa: ")
    _bullet(doc, "Mantive plataforma Wirelist (Starcom) de documentação técnica para montadoras. Spring Boot 2.2, MySQL, Elastic Beanstalk, Angular 15.", "Automotivo/industrial: ")
    _bullet(doc, "Desenvolvi sistema RRZ de gestão documental com assinatura, endosso e auditoria (Hibernate Envers). Spring Boot 3, Java 17, Angular 16.", "Compliance: ")
    _bullet(doc, "Construí app de quizzes Wikle em Go com Gin (backend) e React Native/Expo (frontend).", "Educação: ")
    _bullet(doc, "Participei do ciclo completo de 7 sistemas em 3,5 anos: coleta de requisitos com o cliente, contratos de API em Swagger, UML das classes de domínio, desenvolvimento backend (Java/Spring e Go) e frontend (Angular, React, Vue), integrações com APIs externas e S3, Apache POI para planilhas, SQL nativo quando JPQL não resolvia, testes com JUnit e Mockito, deploy com Docker e Portainer, suporte a aplicações em produção com diagnóstico de defeitos.", "Transversais: ")

    # ---- Formacao ----
    _h2(doc, "Formação")
    _linha_data(doc, "Bacharelado em Ciência da Computação: UFPA", "Abr 2017 - Jan 2024")
    _paragrafo(doc, "Conclusão em janeiro de 2024, colação de grau em março de 2024. Diploma emitido em julho de 2024 (Belém/PA), registro nº 3.209, Livro ICEN-01/24, Folha 29. Concluído em paralelo à atuação profissional iniciada em 2021 e ao período pandêmico, o que estendeu o ciclo total.")

    # ---- Formacao complementar (cursos essenciais) ----
    _h2(doc, "Formação Complementar")
    _bullet(doc, "Udemy (jun/2021). Java Completo Programação Orientada a Objetos.", "Java: ")
    _bullet(doc, "Udemy (out/2021). Microservices do 0 com Spring Cloud, Spring Boot e Docker. Feign, Eureka, API Gateway, Circuit Breaker, Resilience4j.", "Spring/Cloud: ")
    _bullet(doc, "Dev Eficiente (2025, em andamento). Jornada Dev Eficiente: DDD, system design, escalabilidade, CDD, resiliência, testes.", "Arquitetura: ")
    _bullet(doc, "Tech Leads Club (2026, em andamento). Context Engineering, Skills, MCPs, padrões de Subagents e Multi-Agents.", "Tech leadership: ")
    _bullet(doc, "Joshua Bloch. Estudo com anotações em Notion por capítulo: criação de objetos, métodos comuns, classes e interfaces.", "Effective Java: ")
    _bullet(doc, "Casa do Código. OAuth 2.0.", "Segurança: ")
    _bullet(doc, "IGTI/XP (mai/2022). Vue.js, Angular, React, Svelte.", "Frontend: ")
    _bullet(doc, "Udemy (mai/2024, em andamento). Go (Golang): Explorando a Linguagem do Google.", "Go: ")

    # ---- Idiomas ----
    _h2(doc, "Idiomas")
    _bullet(doc, "Nativo.", "Português: ")
    _bullet(doc, "Leitura técnica fluente para documentação, papers e issue trackers. Curso em andamento (2025).", "Inglês: ")

    # ---- IA ----
    _h2(doc, "IA como Eixo de Estudo e Aplicação")
    _bullet(doc, "Benchmark próprio de modelos de embedding para RAG em português e inglês, com corpus desenhado à mão (40 queries parafraseadas, 45 documentos com hard negatives de propósito) e métricas de retrieval (MRR, Recall@K, nDCG@10) em múltiplas execuções. Comparação de 7 modelos. Conclusão documentada: e5-small superou qwen3 em MRR com 2,7x menos espaço e 14x mais velocidade.", "RAG com rigor experimental: ")
    _bullet(doc, "Estudo e aplicação de Context Engineering (Spec Driven, RPI, Rules, Skills, MCPs), padrões de Subagents e Multi-Agents. Estudo comparativo entre revisão solo versus plano e execução com subagents, com métricas, oferecendo evidência preliminar sobre quando delegar a subagents.", "Engenharia agentic: ")
    _bullet(doc, "Prática diária com loop de feedback curto, TDD científico (caso falhando pelo motivo certo, fix mínimo, revert para confirmar red de novo, baby steps) e revisão humana do diff antes de commitar. Material consolidado em notas próprias (Notion e repositórios de estudo).", "Desenvolvimento assistido por IA: ")

    return doc


# ---- Assertions TDD inline ----
def validar(doc: Document) -> None:
    """Valida restricoes ATS no Document gerado. Falha ruidosamente se quebrar."""
    texto = "\n".join(p.text for p in doc.paragraphs)

    # 1. Arquivo abre sem erro (implicito: chamada construir() chegou aqui).
    # 2. Estrutura: secoes obrigatorias presentes.
    secoes = ["PERFIL", "HABILIDADES", "EXPERIÊNCIA", "FORMAÇÃO", "FORMAÇÃO COMPLEMENTAR", "IDIOMAS", "IA COMO EIXO"]
    for s in secoes:
        assert s in texto, f"Secao obrigatoria ausente: {s}"

    # 3. Sem tabelas de layout (ATS nao parseia bem).
    assert len(doc.tables) == 0, f"ATS proibe tabelas de layout. Encontradas: {len(doc.tables)}"

    # 4. Keywords tecnicas para match ATS.
    for kw in ["Java", "Spring Boot", "PostgreSQL", "Clean Architecture", "CQRS", "UFPA"]:
        assert kw in texto, f"Keyword ATS ausente: {kw}"

    # 5. Contato presente.
    assert "(00) 00000-0000" in texto, "Telefone ausente"
    assert "candidato.exemplo@dominio.com" in texto, "E-mail ausente"

    # 6. Regra de honestidade (regra 4 do handoff): jOOQ so no Live2U.
    #    Verifica que jOOQ aparece (Live2U) mas nao no mesmo bullet que Consol/Apontamento.
    assert "jOOQ" in texto, "jOOQ deve aparecer (Live2U)"
    for p in doc.paragraphs:
        if "Consol:" in p.text or "Apontamento:" in p.text:
            assert "jOOQ" not in p.text, f"jOOQ nao deve aparecer em {p.text[:40]}..."

    # 7. Sem em-dashes (regra de estilo do AGENTS.md).
    assert "—" not in texto, "Em-dash encontrado. Usar pontos, virgulas ou dois-pontos."
    assert "–" not in texto, "En-dash encontrado. Usar hifen simples com espacos."

    # 8. Verbo de acao no inicio de cada bullet de experiencia.
    verbos_aceitos = {
        "Desenvolvi", "Modelei", "Integrei", "Mantive", "Colaborei",
        "Construí", "Implementei", "Estruturei", "Participei",
    }
    prefixos_experiencia = {
        "Consol:", "Apontamento:", "Live2U:",
        "Agronegócio:", "Fintech:", "Logística:", "Gamificação corporativa:",
        "Automotivo/industrial:", "Compliance:", "Educação:", "Transversais:",
    }
    for p in doc.paragraphs:
        for prefixo in prefixos_experiencia:
            if p.text.startswith(prefixo):
                corpo = p.text[len(prefixo):].strip()
                # Consol tem "case principal. <verbo>": pular marcadores ate achar o verbo.
                primeira = corpo.split()[0].rstrip(",.") if corpo else ""
                # Se houver marcador antes do verbo (ex.: "case principal."), aceitar.
                palavras = corpo.replace(".", " ").split()
                tem_verbo = any(w.rstrip(",.;:") in verbos_aceitos for w in palavras[:3])
                assert tem_verbo, (
                    f"Bullet '{prefixo}' deve comecar com verbo de acao no passado "
                    f"(entre os 3 primeiros tokens). Inicio: '{primeira}'."
                )

    print("[OK] Validacao ATS passou.")
    print(f"  - Paragrafos: {len(doc.paragraphs)}")
    print(f"  - Tabelas: {len(doc.tables)} (deve ser 0)")
    print(f"  - Caracteres: {len(texto)}")


def main() -> None:
    doc = construir()
    validar(doc)
    doc.save(str(OUTPUT))
    print(f"[OK] Salvo em: {OUTPUT}")


if __name__ == "__main__":
    main()
