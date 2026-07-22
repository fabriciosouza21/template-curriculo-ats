#!/usr/bin/env python3
"""Gera reply_curriculo_java_backend.docx para Candidato Nome.

Variante especifica para vaga "Desenvolvedor(a) Java Pleno" da Reply
Logistic (multinacional italiana de TI, atuacao remota no Brasil). Foco da
JD: produto de Supply Chain. Requisitos: Java backend, microservicos,
Spring/Hibernate, bancos relacionais (MySQL/Oracle), Git/Bitbucket,
superior completo. Diferenciais: DevOps/CI/CD, Docker/Kubernetes, REST,
cloud (AWS/OCI), testes automatizados, mobile (Ionic/Kotlin).

ATS-strict: uma coluna, sem tabelas, sem foto. Fonte 10pt, margens 1.2cm,
line spacing 1.0. Sem em-dash/en-dash.

Conteudo canonico vem de briefing_llm_externo.md. Nada inventado.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Cm, Mm, RGBColor

OUTPUT = Path(__file__).parent / "reply_curriculo_java_backend.docx"

FONTE = "Calibri"
FONTE_NOME = 16
FONTE_H1 = 12
FONTE_H2 = 11
FONTE_CORPO = 10

# Flag de cor: True = layout de uma coluna com cor corporativa. False = ATS
# preto/branco. Padrão das variantes canônicas (Innovo/Zup) é USE_COR=True.
USE_COR = True

# Verde floresta ~#2E8B57. Profissional, ATS-safe em parsers modernos.
# Mesmo verde validado na variante Zup.
COR_DESTAQUE = RGBColor(0x2E, 0x8B, 0x57)


def _style(doc: Document) -> None:
    """A4, margens 1.2cm (compromisso ATS) e fonte padrao."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(FONTE_CORPO)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


def _cor_ambiente() -> RGBColor | None:
    """Retorna a cor de destaque se USE_COR ligado, senao None."""
    return COR_DESTAQUE if USE_COR else None


def _add_borda_inferior(paragraph, cor_hex: str = "2E8B57", tamanho: str = "6") -> None:
    """Adiciona borda inferior ao paragrafo via XML direto."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), tamanho)  # 1/8 pt. 6 = 0.75pt.
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), cor_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _nome(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_NOME)
    r.font.name = FONTE
    cor = _cor_ambiente()
    if cor is not None:
        r.font.color.rgb = cor


def _cargo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_H1)
    r.font.name = FONTE
    cor = _cor_ambiente()
    if cor is not None:
        r.font.color.rgb = cor


def _contato(doc: Document, partes: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(" | ".join(partes))
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _h2(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_H2)
    r.font.name = FONTE
    if USE_COR:
        cor = _cor_ambiente()
        if cor is not None:
            r.font.color.rgb = cor
        _add_borda_inferior(p)


def _paragrafo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _bullet(doc: Document, texto: str, negrito_prefixo: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    if negrito_prefixo:
        r1 = p.add_run(negrito_prefixo)
        r1.bold = True
        r1.font.size = Pt(FONTE_CORPO)
        r1.font.name = FONTE
    r2 = p.add_run(texto)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def _linha_data(doc: Document, esquerda: str, direita: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(18.6), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(esquerda)
    r1.bold = True
    r1.font.size = Pt(FONTE_CORPO)
    r1.font.name = FONTE
    r2 = p.add_run("\t" + direita)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def construir() -> Document:
    doc = Document()
    _style(doc)

    # ---- Cabecalho ----
    _nome(doc, "Candidato Nome")
    _cargo(doc, "Desenvolvedor Backend Java Pleno")
    _contato(
        doc,
        [
            "(00) 00000-0000",
            "candidato.exemplo@dominio.com",
            "Cidade, UF",
            "linkedin.com/in/seu-perfil/",
            "github.com/seu-usuario",
            "seu-portfolio.com/portfolio/seu-usuario",
        ],
    )

    # ---- Perfil ----
    _h2(doc, "Perfil")
    _paragrafo(
        doc,
        "Desenvolvedor Backend Java Pleno com 5 anos de experiencia em "
        "aplicacoes complexas com Java 17 e 21, Spring Boot e Hibernate, em "
        "microservicos e sistemas distribuidos. Passagem por sete dominios "
        "de negocio, com atuacao direta em supply chain e logistica "
        "(marketplace de frete com cotacao, contratacao, notas fiscais e "
        "mensageria). Trabalho com bancos relacionais (MySQL, PostgreSQL), "
        "APIs REST (Representational State Transfer), Git e Bitbucket, "
        "Docker, CI/CD em GitLab e Jenkins, testes automatizados (JUnit, "
        "Mockito, Testcontainers) e nuvem AWS. Referencia tecnica no time em "
        "code review, modelagem de dominios e decisoes de arquitetura "
        "(Clean Architecture, DDD, CQRS). Bacharel em Ciencia da "
        "Computacao pela UFPA (2024).",
    )

    # ---- Habilidades ----
    _h2(doc, "Habilidades")
    _bullet(
        doc,
        "Java (11/16/17/21), Spring Boot (2.x e 3.x), Spring Data JPA, "
        "Spring Security (JWT, OAuth2, OpenID Connect), Spring Cloud (Eureka, "
        "Feign, Resilience4j), Spring WebFlux, Hibernate/JPA.",
        "Backend (JVM): ",
    )
    _bullet(
        doc,
        "REST, OpenAPI/Swagger, SOAP/WSDL, GraphQL (cliente), WebClient.",
        "APIs e contratos: ",
    )
    _bullet(
        doc,
        "Microservicos, sistemas distribuidos, Clean Architecture, CQRS-lite, "
        "arquitetura hexagonal, multi-tenant database-driven, C4 Model.",
        "Arquitetura e sistemas distribuidos: ",
    )
    _bullet(
        doc,
        "PostgreSQL, MySQL, Hibernate/JPA, Hibernate Envers, hibernate-spatial "
        "+ JTS, Flyway, QueryDSL, jOOQ (Live2U), SQL nativo. NoSQL: Redis "
        "(cache, locks, pub/sub), Valkey. Oracle por estudo (curso IGTI 2022).",
        "Bancos relacionais e NoSQL: ",
    )
    _bullet(
        doc,
        "Apache Camel, AWS SQS, WebSocket (STOMP/SockJS). Kafka, RabbitMQ e "
        "Redis Streams em estudo aprofundado.",
        "Mensageria e Event-Driven: ",
    )
    _bullet(
        doc,
        "AWS (ECS, ECR, S3, SQS, SES, CloudFront, Cognito, Lambda, Elastic "
        "Beanstalk, RDS), Docker, Terraform, Git (GitLab, Bitbucket), "
        "GitLab CI/CD, Jenkins. Kubernetes em estudo. Observabilidade: "
        "Sentry, Logstash, OpenTelemetry.",
        "Cloud, DevOps, CI/CD e observabilidade: ",
    )
    _bullet(
        doc,
        "JUnit, Mockito, AssertJ, Testcontainers. Testes de integracao "
        "(@SpringBootTest), seguranca IDOR e concorrencia.",
        "Testes automatizados: ",
    )
    _bullet(
        doc,
        "Design Patterns (Strategy, Chain of Responsibility, Template Method, "
        "State, Command), SOLID, Clean Code, DDD estrategico e tatico.",
        "Boas praticas: ",
    )
    _bullet(
        doc,
        "Kanban e Scrum (Redmine, ClickUp).",
        "Metodologias ageis: ",
    )
    _bullet(
        doc,
        "React Native/Expo e Flutter (mobile). Angular, React e Vue (web) "
        "quando necessario.",
        "Mobile e frontend (complementar): ",
    )

    # ---- Soft Skills e Idiomas ----
    _h2(doc, "Soft Skills e Idiomas")
    _bullet(
        doc,
        "Comunicacao tecnica estruturada a pares: coleta de requisitos, "
        "traducao de problemas de negocio em contratos de API e UML de "
        "dominio, code review e disseminacao de boas praticas.",
    )
    _bullet(
        doc,
        "Pensamento critico na modelagem de dominios complexos e cultura de "
        "testes com Testcontainers, AssertJ e Mockito.",
    )
    _bullet(
        doc,
        "Portugues nativo. Ingles: leitura tecnica fluente, curso em andamento "
        "(2025). Disponibilidade para atuacao remota em equipe global.",
        "Idiomas: ",
    )

    # ---- Experiencia ----
    _h2(doc, "Experiencia")

    # iUsecase
    _linha_data(doc, "iUsecase Tecnologia e Inovacao", "Jul 2025 - Atual")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(
        "Desenvolvedor Backend Pleno 1 (cargo CTPS, CBO 2124-05). Atuacao "
        "remota em tres produtos com foco em design de codigo, arquitetura "
        "limpa, code review e decisoes tecnicas no time."
    )
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(
        doc,
        "Liderei modernizacao de sistema legado de fiscalizacao rodoviaria "
        "para Clean Architecture com CQRS-lite (Commands via ports, leitura "
        "por QueryServices @Cacheable), Java 21 e Spring Boot 3.2. Otimizei "
        "queries com SQL nativo PostgreSQL (window functions, full-text, "
        "UPSERT atomico) e isolei sequenciais com locks pessimistas. "
        "Adicionei testes de seguranca IDOR e de concorrencia, com cobertura "
        "ampla via Testcontainers, AssertJ e Mockito. Cache distribuido em "
        "Redis para sincronizacao offline-first do mobile.",
        "Consol: ",
    )
    _bullet(
        doc,
        "Modelei arquitetura multi-tenant database-driven para timesheet com "
        "policies configuraveis em banco e resolvers @Cacheable em Caffeine, "
        "em arquitetura hexagonal documentada em C4. Spring Boot 3.2 com "
        "Java 17, Spring Security e QueryDSL sobre PostgreSQL. Migracao para "
        "eliminar condicionais de cliente do core em andamento (timelogging "
        "migrado, demais modulos pendentes).",
        "Apontamento: ",
    )
    _bullet(
        doc,
        "Integrei servico externo de IA sobre exames em sistema de saude em "
        "arquitetura hexagonal, com upload de PDF, jobs assincronos (@Async, "
        "Quartz) e eventos pos-commit (@TransactionalEventListener "
        "AFTER_COMMIT). Backend Spring Boot 3.2 com Java 21 e jOOQ para "
        "dashboards. O backend de IA e operado por terceiros; meu trabalho "
        "foi a orquestracao e integracao robusta entre sistemas.",
        "Live2U: ",
    )
    _bullet(
        doc,
        "Mantive pipelines de CI/CD em GitLab com deploy de imagens ARM64 "
        "para ECS via ECR. Colaborei em cultura de testes (Testcontainers, "
        "AssertJ, Mockito), code review e suporte a producao.",
        "Atividades transversais: ",
    )

    # itexto
    _linha_data(doc, "itexto Consultoria em Tecnologia", "Out 2021 - Abr 2025")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(
        "Programador (cargo CTPS, CBO 3171-10). Funcao Full Stack no ciclo "
        "completo de software em sete dominios de negocio (seis em Java e um "
        "em Go)."
    )
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(
        doc,
        "Construi plataforma Ativus de credito rural multi-tenant em Spring "
        "Boot 2.3 e Java 11, com separacao Command/Query em agregados DDD e "
        "SQL nativo para stored procedures multi-schema em PostgreSQL. "
        "Orquestrei integracoes assincronas e gateway de CNDs governamentais "
        "via Apache Camel e AWS SQS. Barter agricola da Corteva em "
        "microsservicos Spring Cloud (Eureka, Feign) e Keycloak OIDC. "
        "Testes com JUnit e Mockito.",
        "Agronegocio: ",
    )
    _bullet(
        doc,
        "Implementei tokenizadora de ativos (QR-Capital) com pipeline "
        "transacional auditavel baseado em Transaction Log/Outbox (unidades "
        "REQUIRES_NEW + REPEATABLE_READ). Cliente GraphQL via "
        "graphql-java-generator + WebClient, com Spring Data JPA, Hibernate "
        "Envers sobre PostgreSQL e OAuth2/Keycloak. Orquestracao assincrona "
        "em Apache Camel + SQS para sync de carteiras (Event-Driven).",
        "Fintech: ",
    )
    _bullet(
        doc,
        "Construi marketplace de frete Flex-Frete (supply chain) com "
        "cotacao, contratacao, notas fiscais e mensageria em Apache Camel "
        "+ SQS. Spring Boot 2.5, PostgreSQL, Redis e React 17. "
        "TransactionTemplate programatico para transacoes de carteira e "
        "cron job de expiracao de solicitacoes.",
        "Logistica e supply chain: ",
    )
    _bullet(
        doc,
        "Estruturei monorepo Weex (bwell) com microservicos async em "
        "Apache Camel + SQS e geracao de certificados em AWS Lambda com "
        "Terraform (gamificacao corporativa).",
        "Gamificacao corporativa: ",
    )
    _bullet(
        doc,
        "Participei do ciclo completo de sete sistemas em tres anos e meio "
        "(incluindo RRZ de compliance com Hibernate Envers e Wirelist para "
        "montadoras): requisitos, contratos de API em Swagger, UML de "
        "dominio, backend Java/Spring, integracoes com APIs externas e S3, "
        "Apache POI, SQL nativo, testes com JUnit e Mockito, deploy com "
        "Docker.",
        "Transversais: ",
    )

    # ---- Formacao ----
    _h2(doc, "Formacao")
    _linha_data(doc, "Bacharelado em Ciencia da Computacao: UFPA", "Abr 2017 - Jan 2024")
    _paragrafo(
        doc,
        "Concluido em janeiro de 2024. Diploma emitido em julho de 2024 "
        "(Belem/PA). Concluido em paralelo a atuacao profissional iniciada "
        "em 2021 e ao periodo pandemico.",
    )

    # ---- Formacao complementar ----
    _h2(doc, "Formacao Complementar")
    _bullet(
        doc,
        "Udemy (out/2021). Microservices do 0 com Spring Cloud, Spring Boot "
        "e Docker.",
        "Microservicos: ",
    )
    _bullet(
        doc,
        "IGTI/XP (set/2022). Analise de Banco de Dados: PostgreSQL, SQL "
        "Server e Oracle.",
        "Bancos relacionais: ",
    )
    _bullet(
        doc,
        "Livro (2025, em andamento). Apache Kafka e Spring Boot.",
        "Mensageria Kafka: ",
    )
    _bullet(
        doc,
        "Dev Eficiente (2025, em andamento). Jornada Dev Eficiente: DDD, "
        "system design, escalabilidade e resiliencia.",
        "Arquitetura e DDD: ",
    )
    _bullet(
        doc,
        "Livros (2024). Desbravando SOLID em Java moderno. Effective Java "
        "(Joshua Bloch).",
        "SOLID e Effective Java: ",
    )
    _bullet(
        doc,
        "Casa do Codigo. OAuth 2.0: fluxos, tokens e escopos.",
        "Seguranca de APIs: ",
    )

    return doc


# ---- Assertions TDD inline ----
def validar(doc: Document) -> None:
    """Valida restricoes ATS e cobertura de keywords da vaga Reply."""
    texto = "\n".join(p.text for p in doc.paragraphs)

    secoes = [
        "PERFIL", "HABILIDADES", "SOFT SKILLS E IDIOMAS",
        "EXPERIENCIA", "FORMACAO", "FORMACAO COMPLEMENTAR",
    ]
    for s in secoes:
        assert s in texto, f"Secao obrigatoria ausente: {s}"

    assert "Idiomas:" in texto, "Bullet de Idiomas ausente"

    assert len(doc.tables) == 0, f"ATS proibe tabelas. Encontradas: {len(doc.tables)}"

    keywords_vaga = [
        "Java", "Java 17", "Java 21", "Spring Boot", "Spring Data JPA",
        "Spring Security", "JPA", "Hibernate",
        "REST", "microservicos", "sistemas distribuidos", "API",
        "PostgreSQL", "MySQL", "Oracle",
        "Apache Camel", "AWS SQS",
        "AWS", "Git", "Bitbucket", "CI/CD", "Docker", "Kubernetes",
        "JUnit", "Mockito", "Testcontainers",
        "code review", "nuvem",
        "supply chain",
        "SQL nativo",
    ]
    for kw in keywords_vaga:
        assert kw in texto, f"Keyword da vaga ausente: {kw}"

    proibidos_precisao = [
        "32 queries", "16 repositorios", "7 testes", "TTL 1h",
        "123 commands", "111 queries", "13 metodos", "10 cron",
        "720 classes", "537", "629 classes", "14 rotas", "5 unidades",
        "3 locks",
    ]
    for termo in proibidos_precisao:
        assert termo not in texto, f"Metrica de falsa precisao (cortar): {termo}"

    proibidos = [
        "WireMock em producao",
        "Kubernetes em producao",
        "Grafana em producao",
        "Micronaut em producao",
        "Micronaut",
        "Oracle em producao",
        "OCI em producao",
        "Ionic",
        "Kotlin",
        "experiencia em SAP",
    ]
    for termo in proibidos:
        assert termo not in texto, f"Termo proibido (sem evidencia): {termo}"

    siglas_expandidas = [
        "Representational State Transfer",
        "OpenID Connect",
    ]
    for sigla in siglas_expandidas:
        assert sigla in texto, f"Sigla nao expandida: {sigla}"

    assert "(00) 00000-0000" in texto, "Telefone ausente"
    assert "candidato.exemplo@dominio.com" in texto, "E-mail ausente"

    assert "Desenvolvedor Backend Pleno 1" in texto, "Cargo CTPS iUsecase ausente"
    assert "CBO 2124-05" in texto, "CBO iUsecase ausente"
    assert "CBO 3171-10" in texto, "CBO itexto ausente"

    assert "jOOQ" in texto, "jOOQ deve aparecer (Live2U)"
    for p in doc.paragraphs:
        if "Consol:" in p.text or "Apontamento:" in p.text:
            assert "jOOQ" not in p.text, f"jOOQ nao deve aparecer em {p.text[:40]}..."

    assert "—" not in texto, "Em-dash encontrado."
    assert "–" not in texto, "En-dash encontrado."

    verbos_aceitos = {
        "Desenvolvi", "Modelei", "Integrei", "Mantive", "Colaborei",
        "Construi", "Implementei", "Estruturei", "Participei",
        "Liderei", "Otimizei", "Adicionei", "Orquestrei", "Usei",
    }
    prefixos_experiencia = {
        "Consol:", "Apontamento:", "Live2U:", "Atividades transversais:",
        "Agronegocio:", "Fintech:", "Logistica e supply chain:",
        "Gamificacao corporativa:",
        "Automotivo/industrial:", "Compliance:", "Transversais:",
    }
    for p in doc.paragraphs:
        for prefixo in prefixos_experiencia:
            if p.text.startswith(prefixo):
                corpo = p.text[len(prefixo):].strip()
                palavras = corpo.replace(".", " ").split()
                tem_verbo = any(w.rstrip(",.;:") in verbos_aceitos for w in palavras[:3])
                assert tem_verbo, (
                    f"Bullet '{prefixo}' deve comecar com verbo de acao no "
                    f"passado (entre os 3 primeiros tokens). Inicio: "
                    f"'{palavras[0] if palavras else ''}'."
                )

    print("[OK] Validacao ATS + cobertura de keywords da vaga Reply passou.")
    print(f"  - Paragrafos: {len(doc.paragraphs)}")
    print(f"  - Tabelas: {len(doc.tables)} (deve ser 0)")
    print(f"  - Caracteres: {len(texto)}")


def main() -> None:
    doc = construir()
    validar(doc)
    doc.save(str(OUTPUT))
    print(f"[OK] Salvo em: {OUTPUT}")


if __name__ == "__main__":
    main()
