#!/usr/bin/env python3
"""Gera innvo_curriculo_java_senior.docx para Candidato Nome.

Variante especifica para vaga "Desenvolvedor(a) Java Senior" da Innvo Labs.
Foco da JD: Java 17+, Spring (Boot/MVC/Data/Security/JPA), APIs REST,
microsservicos, bancos relacionais (PostgreSQL/Oracle/MySQL) e NoSQL
(MongoDB/Redis), mensageria (Kafka/RabbitMQ), Docker, Kubernetes, CI/CD,
testes automatizados (JUnit/Mockito), Design Patterns, SOLID, Clean Code,
metodologias ageis. Diferenciais: Event-Driven, DDD, Clean Architecture,
alta disponibilidade, cloud (AWS/Azure/GCP), observabilidade (Prometheus,
Grafana, ELK, OpenTelemetry, Dynatrace).

Diferencas vs zup_curriculo_backend.py:
- Perfil reescrito com keywords Senior (arquitetura, microsservicos,
  disseminacao de boas praticas, code review, escalabilidade).
- Habilidades reordenadas com estrelas sendo Spring (Boot/MVC/Data/
  Security), arquitetura de microsservicos, mensageria, design de codigo
  (SOLID/Clean Code/Design Patterns) e observabilidade.
- Experiencia elevada a tom de referencia tecnica (code review, decisoes
  de arquitetura, disseminacao de boas praticas) com honestidade sobre o
  que e producao x estudo (Kafka estudo, K8s nao afirmado como prod).
- Frontend recolhido a uma linha. Mobile e Go removidos.
- Secao IA removida (vaga nao pede). Cursos priorizam Spring, arquitetura
  (DDD, Clean Architecture) e mensageria.

ATS-strict: uma coluna, sem tabelas, sem cores, sem foto. Fonte 10pt,
margens 1.2cm, line spacing 1.05. Sem em-dash/en-dash.

Conteudo canonico vem de briefing_llm_externo.md. Nada inventado.
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Cm, Mm, RGBColor

OUTPUT = Path(__file__).parent / "innvo_curriculo_java_senior.docx"

FONTE = "Calibri"
FONTE_NOME = 16
FONTE_H1 = 12
FONTE_H2 = 11
FONTE_CORPO = 10

# Flag de cor: True = layout verde de uma coluna. False = ATS preto/branco.
USE_COR = True

# Azul corporativo ~#1F4E79. Profissional, alinhado a vaga senior,
# ATS-safe em parsers modernos.
COR_DESTAQUE = RGBColor(0x1F, 0x4E, 0x79)


def _style(doc: Document) -> None:
    """A4, margens 1.2cm (compromisso ATS) e fonte padrao."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(1.2)
        section.bottom_margin = Cm(1.2)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(FONTE_CORPO)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


def _cor_ambiente() -> RGBColor | None:
    """Retorna a cor de destaque se USE_COR ligado, senao None."""
    return COR_DESTAQUE if USE_COR else None


def _nome(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_NOME)
    r.font.name = FONTE
    cor = _cor_ambiente()
    if cor is not None:
        r.font.color.rgb = cor


def _cargo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_H1)
    r.font.name = FONTE
    cor = _cor_ambiente()
    if cor is not None:
        r.font.color.rgb = cor


def _contato(doc: Document, partes: list[str]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(" | ".join(partes))
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _add_borda_inferior(paragraph, cor_hex: str = "1F4E79", tamanho: str = "6") -> None:
    """Adiciona borda inferior ao paragrafo via XML direto."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), tamanho)  # 1/8 pt. 6 = 0.75pt.
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), cor_hex)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _h2(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_H2)
    r.font.name = FONTE
    if USE_COR:
        cor = _cor_ambiente()
        if cor is not None:
            r.font.color.rgb = cor
        _add_borda_inferior(p)


def _paragrafo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def _bullet(doc: Document, texto: str, negrito_prefixo: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.5)
    if negrito_prefixo:
        r1 = p.add_run(negrito_prefixo)
        r1.bold = True
        r1.font.size = Pt(FONTE_CORPO)
        r1.font.name = FONTE
    r2 = p.add_run(texto)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def _linha_data(doc: Document, esquerda: str, direita: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(18.6), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(esquerda)
    r1.bold = True
    r1.font.size = Pt(FONTE_CORPO)
    r1.font.name = FONTE
    r2 = p.add_run("\t" + direita)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def construir() -> Document:
    doc = Document()
    _style(doc)

    # ---- Cabecalho ----
    _nome(doc, "Candidato Nome")
    _cargo(doc, "Desenvolvedor Java Senior")
    _contato(
        doc,
        [
            "(00) 00000-0000",
            "candidato.exemplo@dominio.com",
            "Cidade, UF",
            "linkedin.com/in/seu-perfil/",
            "github.com/seu-usuario",
            "seu-portfolio.com/portfolio/seu-usuario",
        ],
    )

    # ---- Perfil (focado nas keywords Senior da vaga) ----
    _h2(doc, "Perfil")
    _paragrafo(
        doc,
        "Desenvolvedor Java com 5 anos de experiencia em arquitetura de software, "
        "microsservicos Spring Boot e APIs REST (Representational State Transfer), "
        "com passagem por sete dominios de negocio distintos (agronegocio, fintech, "
        "logistica, automotivo, gamificacao corporativa, compliance e educacao). "
        "Atuo como referencia tecnica no time em code review, disseminacao de boas "
        "praticas (SOLID, Clean Code, Design Patterns), modelagem de dominios "
        "complexos e decisoes de arquitetura (Clean Architecture, DDD, CQRS). "
        "Experiencia com integracao entre sistemas via Apache Camel e mensageria, "
        "bancos relacionais e NoSQL (Redis, MongoDB em estudo aplicado), pipelines "
        "de CI/CD em GitLab e Jenkins, e observabilidade de aplicacoes em producao "
        "com Sentry, Logstash e OpenTelemetry. Bacharel em Ciencia da Computacao "
        "pela UFPA (2024).",
    )

    # ---- Habilidades (reordenadas para a vaga) ----
    _h2(doc, "Habilidades")
    _bullet(
        doc,
        "Java (11/16/17/21), Spring Boot (2.x e 3.x), Spring MVC, Spring Data JPA, "
        "Spring Security (JWT, OAuth2, OpenID Connect via Keycloak e AWS Cognito), "
        "Spring Cloud (Eureka, Feign, API Gateway, Circuit Breaker, Resilience4j), "
        "Spring WebFlux.",
        "Backend (JVM): ",
    )
    _bullet(
        doc,
        "REST (Representational State Transfer), OpenAPI/Swagger, SOAP/WSDL "
        "(integracao legada), GraphQL (cliente via graphql-java-generator), WebClient.",
        "APIs e contratos: ",
    )
    _bullet(
        doc,
        "Arquitetura de microsservicos, sistemas distribuidos, Clean Architecture, "
        "CQRS-lite (Commands via ports, leitura por QueryServices), arquitetura "
        "hexagonal, multi-tenant database-driven, documentacao em C4 Model.",
        "Arquitetura: ",
    )
    _bullet(
        doc,
        "PostgreSQL, MySQL, Oracle (curso IGTI/XP), Hibernate/JPA, Hibernate Envers "
        "(auditoria), hibernate-spatial + JTS (geometria), Flyway (migrations), "
        "QueryDSL, jOOQ (Live2U), SQL nativo para otimizacao de performance "
        "(window functions, full-text com unaccent PT-BR, UPSERT atomico).",
        "Bancos relacionais: ",
    )
    _bullet(
        doc,
        "Redis (cache, locks distribuidos, pub/sub), Valkey. MongoDB em estudo "
        "aplicado (estudo autoral em Notion).",
        "Bancos NoSQL: ",
    )
    _bullet(
        doc,
        "Apache Camel (roteamento B2B, gateway de CNDs governamentais), AWS SQS, "
        "WebSocket (STOMP/SockJS). Kafka em estudo aprofundado (topicos, parties, "
        "consumer groups, offsets, comparativo Kafka vs SQS/SNS), RabbitMQ, Redis "
        "Streams como alternativas de message broker.",
        "Mensageria e Event-Driven: ",
    )
    _bullet(
        doc,
        "Design Patterns (Strategy, Chain of Responsibility, Template Method, State, "
        "Command, Observer, Adapter), SOLID, Clean Code, modular architecture, "
        "metricas de coesao e acoplamento, DDD estrategico e tatico (aggregates, "
        "value objects, use cases, domain services).",
        "Boas praticas: ",
    )
    _bullet(
        doc,
        "JUnit, Mockito, AssertJ, Testcontainers, WireMock. Testes de integracao "
        "ponta-a-ponta com @SpringBootTest, testes de seguranca IDOR, testes de "
        "concorrencia em sequenciais. Cobertura ampla por dominio.",
        "Testes automatizados: ",
    )
    _bullet(
        doc,
        "Git (GitFlow, trunk-based), GitLab CI/CD, Jenkins, Docker, Portainer, "
        "Terraform. Pipelines de CI/CD com deploy de imagens ARM64 para ECS via "
        "ECR. AWS (ECS, ECR, S3, SQS, SES, CloudFront, Cognito, Lambda, Elastic "
        "Beanstalk, RDS). Kubernetes em estudo (nao em producao).",
        "DevOps e CI/CD: ",
    )
    _bullet(
        doc,
        "Sentry, Logstash, OpenTelemetry para monitoramento de aplicacoes em "
        "producao. Prometheus, Grafana, ELK Stack e Dynatrace em estudo.",
        "Observabilidade: ",
    )
    _bullet(
        doc,
        "Scrum e Kanban (Redmine, ClickUp). Atuacao em equipes multidisciplinares "
        "com coleta de requisitos, planejamento de tarefas e rituais ageis.",
        "Metodologias ageis: ",
    )
    _bullet(
        doc,
        "Angular, React e Vue para atuacao full stack quando necessario.",
        "Frontend (complementar): ",
    )

    # ---- Soft Skills e Idiomas ----
    _h2(doc, "Soft Skills e Idiomas")
    _bullet(
        doc,
        "Comunicacao tecnica estruturada a pares: coleta de requisitos com cliente, "
        "traducao de problemas de negocio em contratos de API e UML de dominio, "
        "disseminacao de boas praticas no time.",
    )
    _bullet(
        doc,
        "Pensamento critico na revisao tecnica de codigo e na modelagem de dominios "
        "complexos (inspecoes rodoviarias, exames de saude, apontamento de horas, "
        "credito rural, tokenizacao de ativos).",
    )
    _bullet(
        doc,
        "Colaboracao em cultura de testes e revisao por pares (Testcontainers, "
        "AssertJ, Mockito), mentor informal em boas praticas de codigo.",
    )
    _bullet(
        doc,
        "Portugues nativo. Ingles: leitura tecnica fluente, curso em andamento (2025).",
        "Idiomas: ",
    )

    # ---- Experiencia ----
    _h2(doc, "Experiencia")

    # iUsecase
    _linha_data(doc, "iUsecase Tecnologia e Inovacao", "Jul 2025 - Atual")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(
        "Desenvolvedor Backend Pleno 1 (cargo CTPS, CBO 2124-05). Atuacao remota "
        "em tres produtos com foco em design de codigo, arquitetura limpa, code "
        "review e decisoes tecnicas no time."
    )
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(
        doc,
        "Liderei modernizacao de sistema legado de fiscalizacao rodoviaria "
        "migrando para Clean Architecture com CQRS-lite (Commands via ports, "
        "leitura por QueryServices @Cacheable), Java 21 e Spring Boot 3.2 com "
        "Spring Data JPA e Spring Security. Otimizei queries criticas com SQL "
        "nativo PostgreSQL (window functions, full-text com unaccent PT-BR, "
        "UPSERT atomico) e isolei sequenciais concorridos com locks pessimistas "
        "(PESSIMISTIC_WRITE). Adicionei testes de seguranca IDOR cobrindo "
        "controllers e teste de concorrencia no sequencial, com cobertura ampla "
        "via Testcontainers, AssertJ e Mockito. Cache distribuido em Redis com "
        "lock para sincronizacao offline-first do mobile.",
        "Consol: ",
    )
    _bullet(
        doc,
        "Modelei arquitetura multi-tenant database-driven para timesheet com "
        "policies configuraveis em banco e resolvers @Cacheable em Caffeine, em "
        "arquitetura hexagonal documentada em C4. Liderei migracao para eliminar "
        "condicionais de cliente do core: modulo de timelogging migrado, demais "
        "modulos ainda pendentes (divida documentada). Spring Boot 3.2 com Java "
        "17, Spring Security e QueryDSL para queries dinamicas tipadas sobre "
        "PostgreSQL.",
        "Apontamento: ",
    )
    _bullet(
        doc,
        "Integrei servico externo de IA sobre exames em sistema de saude em "
        "arquitetura hexagonal, com upload de PDF, jobs assincronos (@Async, "
        "Quartz para refresh de tokens de parceiros) e eventos pos-commit "
        "(@TransactionalEventListener AFTER_COMMIT para WhatsApp e sync de "
        "links). Backend Spring Boot 3.2 com Java 21 e jOOQ para dashboards. O "
        "backend de IA e operado por terceiros; meu trabalho foi a orquestracao "
        "e a integracao robusta entre sistemas distintos.",
        "Live2U: ",
    )
    _bullet(
        doc,
        "Mantive pipelines de CI/CD em GitLab com deploy de imagens ARM64 para "
        "ECS via ECR (backend) e S3 com CloudFront (frontend). Colaborei em "
        "cultura de testes com Testcontainers, AssertJ e Mockito, revisao "
        "tecnica de codigo no time e suporte a aplicacoes em producao com "
        "diagnostico de defeitos.",
        "Atividades transversais: ",
    )

    # itexto
    _linha_data(doc, "itexto Consultoria em Tecnologia", "Out 2021 - Abr 2025")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(
        "Programador (cargo CTPS, CBO 3171-10). Funcao Full Stack no ciclo "
        "completo de software em sete dominios de negocio (seis em Java e um "
        "em Go)."
    )
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE

    _bullet(
        doc,
        "Construi plataforma Ativus de credito rural multi-tenant em Spring Boot "
        "2.3 e Java 11, com separacao Command/Query em agregados DDD e SQL nativo "
        "para stored procedures multi-schema em PostgreSQL. Orquestrei integracoes "
        "assincronas com @Async e cron jobs (Serasa, gateway Vindi, SendGrid) e "
        "gateway de CNDs governamentais via Apache Camel e AWS SQS. Implementei "
        "barter agricola da Corteva em microsservicos Spring Cloud (Eureka, Feign) "
        "e Keycloak OIDC. Cobertura ampla de testes com JUnit e Mockito.",
        "Agronegocio: ",
    )
    _bullet(
        doc,
        "Implementei tokenizadora de ativos (QR-Capital) com pipeline transacional "
        "auditavel baseado em Transaction Log/Outbox (unidades REQUIRES_NEW + "
        "REPEATABLE_READ com rastreabilidade via ThreadLocal). Cliente GraphQL "
        "via graphql-java-generator + WebClient, com Spring Data JPA, Hibernate "
        "Envers sobre PostgreSQL e OAuth2/Keycloak. Orquestracao assincrona em "
        "Apache Camel + SQS para sync de carteiras e status de transferencias "
        "bancarias (Event-Driven na pratica).",
        "Fintech: ",
    )
    _bullet(
        doc,
        "Desenvolvi marketplace de frete Flex-Frete com cotacao, contratacao e "
        "notas fiscais. Usei TransactionTemplate programatico para transacoes "
        "financeiras de carteira (escopo condicional debito/credito) e cron job "
        "de expiracao/cancelamento de fretes. Spring Boot 2.5, PostgreSQL, Redis "
        "(cache), React 17.",
        "Logistica: ",
    )
    _bullet(
        doc,
        "Estruturei monorepo Weex com microsservico async dedicado (weex.async) "
        "com rotas Apache Camel + SQS para certificados, expurgo, logs e "
        "processamento paralelo de imagens (CompletableFuture). Geracao de "
        "certificados em AWS Lambda e IaC com Terraform. Cobertura ampla de "
        "testes de integracao ponta-a-ponta com @SpringBootTest.",
        "Gamificacao corporativa: ",
    )
    _bullet(
        doc,
        "Participei do ciclo completo de sete sistemas em tres anos e meio "
        "(incluindo RRZ de compliance com auditoria via Hibernate Envers e "
        "Wirelist de documentacao tecnica para montadoras): coleta de requisitos "
        "com o cliente, contratos de API em Swagger, UML das classes de dominio, "
        "desenvolvimento backend (Java/Spring), integracoes com APIs externas e "
        "S3, Apache POI para planilhas, SQL nativo para otimizacao de performance "
        "quando JPQL nao resolvia, testes com JUnit e Mockito, deploy com Docker "
        "e Portainer, suporte a aplicacoes em producao com diagnostico de defeitos.",
        "Transversais: ",
    )

    # ---- Formacao ----
    _h2(doc, "Formacao")
    _linha_data(doc, "Bacharelado em Ciencia da Computacao: UFPA", "Abr 2017 - Jan 2024")
    _paragrafo(
        doc,
        "Concluido em janeiro de 2024. Diploma emitido em julho de 2024 "
        "(Belem/PA). Concluido em paralelo a atuacao profissional iniciada em "
        "2021 e ao periodo pandemico.",
    )

    # ---- Formacao complementar (prioriza Spring, arquitetura, mensageria) ----
    _h2(doc, "Formacao Complementar")
    _bullet(
        doc,
        "Udemy (out/2021). Microservices do 0 com Spring Cloud, Spring Boot e "
        "Docker (Feign, Eureka, Circuit Breaker, Resilience4j).",
        "Microsservicos: ",
    )
    _bullet(
        doc,
        "Dev Eficiente (2025, em andamento). Jornada Dev Eficiente: DDD, system "
        "design, escalabilidade, CDD, resiliencia.",
        "Arquitetura e DDD: ",
    )
    _bullet(
        doc,
        "Livro (2025, em andamento). Apache Kafka e Spring Boot.",
        "Mensageria Kafka: ",
    )
    _bullet(
        doc,
        "Livros (2024). Desbravando SOLID em Java moderno. Effective Java "
        "(Joshua Bloch) com anotacoes em Notion por capitulo.",
        "SOLID, Clean Code e Effective Java: ",
    )
    _bullet(
        doc,
        "Casa do Codigo. OAuth 2.0: fluxos de autorizacao, tokens e escopos.",
        "Seguranca de APIs: ",
    )

    return doc


# ---- Assertions TDD inline ----
def validar(doc: Document) -> None:
    """Valida restricoes ATS e cobertura de keywords da vaga Innvo Labs."""
    texto = "\n".join(p.text for p in doc.paragraphs)

    # 1. Estrutura: secoes obrigatorias presentes.
    secoes = [
        "PERFIL", "HABILIDADES", "SOFT SKILLS E IDIOMAS",
        "EXPERIENCIA", "FORMACAO", "FORMACAO COMPLEMENTAR",
    ]
    for s in secoes:
        assert s in texto, f"Secao obrigatoria ausente: {s}"

    # 1b. Idiomas deve aparecer.
    assert "Idiomas:" in texto, "Bullet de Idiomas ausente"

    # 2. Sem tabelas de layout (ATS nao parseia bem).
    assert len(doc.tables) == 0, f"ATS proibe tabelas. Encontradas: {len(doc.tables)}"

    # 3. Keywords da vaga (obrigatorias) para match ATS.
    keywords_vaga = [
        # Backend core
        "Java", "Spring Boot", "Spring MVC", "Spring Data JPA",
        "Spring Security", "JPA", "Hibernate",
        # REST e arquitetura
        "REST", "microsservicos", "API",
        # Bancos
        "PostgreSQL", "MySQL", "Oracle", "Redis", "MongoDB",
        # Mensageria
        "Apache Camel", "Kafka", "RabbitMQ", "AWS SQS",
        # DevOps
        "Git", "CI/CD", "Docker",
        # Testes
        "JUnit", "Mockito",
        # Boas praticas
        "Design Patterns", "SOLID", "Clean Code",
        # Ageis
        "Scrum", "Kanban",
        # Diferenciais
        "DDD", "Clean Architecture", "OpenTelemetry",
        # Literais da JD
        "code review", "boas praticas",
        # Padroes tecnicos reais (defensaveis em entrevista)
        "SQL nativo", "REQUIRES_NEW", "Testcontainers",
        "Outbox", "Event-Driven",
    ]
    for kw in keywords_vaga:
        assert kw in texto, f"Keyword da vaga ausente: {kw}"

    # 3d. Anti-precisao-falsa: proibe contagens inventadas.
    proibidos_precisao = [
        "32 queries", "16 repositorios", "7 testes", "TTL 1h",
        "123 commands", "111 queries", "13 metodos", "10 cron",
        "720 classes", "537", "629 classes", "14 rotas", "5 unidades",
        "3 locks",
    ]
    for termo in proibidos_precisao:
        assert termo not in texto, f"Metrica de falsa precisao (cortar): {termo}"

    # 3c. Anti-invencao: termos que NAO devem aparecer como experiencia em prod.
    # Kafka e estudo, nao producao. K8s nao afirmar como experiencia direta.
    proibidos = [
        "WireMock em producao",  # no pom mas sem uso real
        "Kubernetes em producao",  # nao tenho K8s em prod
        "Dynatrace em producao",  # estudo, nao prod
        "Grafana em producao",  # estudo, nao prod
        "experiencia em SAP",
    ]
    for termo in proibidos:
        assert termo not in texto, f"Termo proibido (sem evidencia): {termo}"

    # 3b. Siglas expandidas ao menos uma vez (regra Alura).
    siglas_expandidas = [
        "Representational State Transfer",
        "OpenID Connect",
    ]
    for sigla in siglas_expandidas:
        assert sigla in texto, f"Sigla nao expandida: {sigla}"

    # 4. Contato presente.
    assert "(00) 00000-0000" in texto, "Telefone ausente"
    assert "candidato.exemplo@dominio.com" in texto, "E-mail ausente"

    # 5. Cargos CTPS explicitados (honestidade).
    assert "Desenvolvedor Backend Pleno 1" in texto, "Cargo CTPS iUsecase ausente"
    assert "CBO 2124-05" in texto, "CBO iUsecase ausente"
    assert "CBO 3171-10" in texto, "CBO itexto ausente"

    # 6. Regra de honestidade: jOOQ so no Live2U.
    assert "jOOQ" in texto, "jOOQ deve aparecer (Live2U)"
    for p in doc.paragraphs:
        if "Consol:" in p.text or "Apontamento:" in p.text:
            assert "jOOQ" not in p.text, f"jOOQ nao deve aparecer em {p.text[:40]}..."

    # 7. Sem em-dashes (regra de estilo do AGENTS.md).
    assert "—" not in texto, "Em-dash encontrado."
    assert "–" not in texto, "En-dash encontrado."

    # 8. Verbo de acao no inicio de cada bullet de experiencia.
    verbos_aceitos = {
        "Desenvolvi", "Modelei", "Integrei", "Mantive", "Colaborei",
        "Construi", "Implementei", "Estruturei", "Participei",
        "Liderei", "Otimizei", "Adicionei", "Orquestrei", "Usei",
    }
    prefixos_experiencia = {
        "Consol:", "Apontamento:", "Live2U:", "Atividades transversais:",
        "Agronegocio:", "Fintech:", "Logistica:", "Gamificacao corporativa:",
        "Automotivo/industrial:", "Compliance:", "Transversais:",
    }
    for p in doc.paragraphs:
        for prefixo in prefixos_experiencia:
            if p.text.startswith(prefixo):
                corpo = p.text[len(prefixo):].strip()
                palavras = corpo.replace(".", " ").split()
                tem_verbo = any(w.rstrip(",.;:") in verbos_aceitos for w in palavras[:3])
                assert tem_verbo, (
                    f"Bullet '{prefixo}' deve comecar com verbo de acao no "
                    f"passado (entre os 3 primeiros tokens). Inicio: "
                    f"'{palavras[0] if palavras else ''}'."
                )

    print("[OK] Validacao ATS + cobertura de keywords da vaga Innvo Labs passou.")
    print(f"  - Paragrafos: {len(doc.paragraphs)}")
    print(f"  - Tabelas: {len(doc.tables)} (deve ser 0)")
    print(f"  - Caracteres: {len(texto)}")


def main() -> None:
    doc = construir()
    validar(doc)
    doc.save(str(OUTPUT))
    print(f"[OK] Salvo em: {OUTPUT}")


if __name__ == "__main__":
    main()
