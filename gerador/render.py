"""Helpers de renderização ATS para python-docx.

A4, margens 2cm, fonte Calibri 10.5pt, espaçamentos controlados.

Estilo ATS único: sem tabelas de layout, uma coluna, texto puro. A cor de
destaque é opcional (default None = preto, ATS-friendly). Quando o manifesto
passa uma cor (ex.: #0B1641), ela é aplicada apenas em nome, cargo e H2,
mantendo o corpo preto. Os nomes são públicos (sem underscore) para consumo
limpo por `gerador.montar`.
"""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.shared import Pt, Cm, Mm, RGBColor

# Fonte ATS-safe: Calibri (fallback universal). Tamanhos alinhados às
# recomendações da Alura (10-12pt) com margens ~2cm. Aceita até 3 páginas
# para perfis com bagagem.
FONTE = "Calibri"
FONTE_NOME = 18
FONTE_H1 = 12
FONTE_H2 = 11
FONTE_CORPO = 10.5


def _hex_para_rgb(hex_str):
    """Converte "#RRGGBB" ou "RRGGBB" em RGBColor. None/vazio -> None.

    Levanta ValueError se a string não casar com 6 hex digits.
    """
    if hex_str is None:
        return None
    s = str(hex_str).strip().lstrip("#")
    if len(s) != 6:
        raise ValueError(f"cor hex inválida (esperado #RRGGBB): {hex_str!r}")
    try:
        r = int(s[0:2], 16)
        g = int(s[2:4], 16)
        b = int(s[4:6], 16)
    except ValueError as exc:
        raise ValueError(f"cor hex inválida (não é hex): {hex_str!r}") from exc
    return RGBColor(r, g, b)


def style(doc: Document) -> None:
    """Configura A4, margens 2cm e fonte Calibri 10.5pt no estilo Normal."""
    for section in doc.sections:
        section.page_height = Mm(297)
        section.page_width = Mm(210)
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = FONTE
    normal.font.size = Pt(FONTE_CORPO)
    normal.paragraph_format.space_after = Pt(1)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.line_spacing = 1.1


def nome(doc: Document, texto: str, cor=None) -> None:
    """Nome em maiúsculas, 18pt bold, centralizado. `cor` opcional."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_NOME)
    r.font.name = FONTE
    if cor is not None:
        r.font.color.rgb = cor


def cargo(doc: Document, texto: str, cor=None) -> None:
    """Cargo logo abaixo do nome, 12pt, centralizado. `cor` opcional."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_H1)
    r.font.name = FONTE
    if cor is not None:
        r.font.color.rgb = cor


def contato(doc: Document, partes: list[str]) -> None:
    """Linha de contato: partes unidas por ' | ', 10.5pt, centralizada."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(" | ".join(partes))
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def h2(doc: Document, texto: str, cor=None) -> None:
    """Cabeçalho de seção em maiúsculas, bold, 11pt. `cor` opcional.

    Sem borda inferior: underline em run separado seria ruído para ATS.
    Mantém apenas bold + caixa alta. A cor, quando fornecida, dá
    identidade visual sem comprometer a leitura ATS.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto.upper())
    r.bold = True
    r.font.size = Pt(FONTE_H2)
    r.font.name = FONTE
    if cor is not None:
        r.font.color.rgb = cor


def paragrafo(doc: Document, texto: str) -> None:
    """Parágrafo de corpo, 10.5pt."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(texto)
    r.font.size = Pt(FONTE_CORPO)
    r.font.name = FONTE


def bullet(doc: Document, texto: str, negrito_prefixo: str = "") -> None:
    """Bullet com prefixo em negrito opcional.

    ATS lê bullet char '-' (estilo List Bullet) bem. Quando
    `negrito_prefixo` é fornecido, gera um run bold seguido de um run
    normal; caso contrário, um único run normal.
    """
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    if negrito_prefixo:
        r1 = p.add_run(negrito_prefixo)
        r1.bold = True
        r1.font.size = Pt(FONTE_CORPO)
        r1.font.name = FONTE
    r2 = p.add_run(texto)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE


def linha_data(doc: Document, esquerda: str, direita: str) -> None:
    """Linha com cargo a esquerda (bold) e data a direita via tab stop.

    Tab stop à direita em 17cm (limite útil da área de escrita com margens
    de 2cm em A4: 21 - 2 - 2 = 17cm).
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Cm(17.0), WD_TAB_ALIGNMENT.RIGHT)
    r1 = p.add_run(esquerda)
    r1.bold = True
    r1.font.size = Pt(FONTE_CORPO)
    r1.font.name = FONTE
    r2 = p.add_run("\t" + direita)
    r2.font.size = Pt(FONTE_CORPO)
    r2.font.name = FONTE
