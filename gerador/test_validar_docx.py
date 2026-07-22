"""Testes da camada de assertions ATS em gerador.validar_docx.

Padrão TDD: testes escritos antes da implementação. Cada teste constrói
um `Document` sintético usando os helpers públicos de `gerador.render`
(nome, cargo, contato, h2, paragrafo, bullet, linha_data) e então chama
`validar(doc)` (ou `validar(doc, manifesto)` quando relevante).

Cobre o contrato público da Task 4 do plano gerador-docx:
- documento válido passa sem levantar;
- tabela presente -> AssertionError mencionando "tabela";
- em-dash (U+2014) presente -> AssertionError mencionando "em-dash";
- seção obrigatória ausente -> AssertionError;
- verbo de ação ausente em bullet de experiência -> AssertionError
  mencionando "verbo";
- consistência manifesto: ia False omitida; case pedido ausente falha.
"""

import pytest
from docx import Document

from gerador.render import (
    style,
    nome,
    cargo,
    contato,
    h2,
    paragrafo,
    bullet,
    linha_data,
)
from gerador.validar_docx import validar


# ---- Fixtures sintéticas ----

# Contato canônico real do brief (telefone + email). Outros campos podem
# variar por teste. Estes dois são exigidos pelas regras de validação.
_TELEFONE = "(00) 00000-0000"
_EMAIL = "candidato.exemplo@dominio.com"


def _doc_valido_minimo() -> Document:
    """Document sintético que satisfaz todas as regras ATS obrigatórias.

    Contém: nome, cargo, contato (telefone + email reais), seções
    PERFIL, HABILIDADES, EXPERIÊNCIA, FORMAÇÃO, e um bullet de
    experiência começando com verbo de ação aceito ("Construí").
    """
    doc = Document()
    style(doc)
    nome(doc, "Candidato Nome")
    cargo(doc, "Engenheiro de Software Pleno")
    contato(doc, [_TELEFONE, _EMAIL, "Belém, PA"])

    h2(doc, "Perfil")
    paragrafo(doc, "Perfil sintético para teste de validação DOCX.")

    h2(doc, "Habilidades")
    bullet(doc, "Java, Spring Boot, PostgreSQL", negrito_prefixo="Backend (JVM): ")

    h2(doc, "Experiência")
    linha_data(doc, "iUsecase", "Jan 2023 - Atual")
    # Corpo começa com "case principal." antes do verbo, replicando o
    # padrão real do YAML (iusecase.yml). A regra aceita verbo nos 3
    # primeiros tokens do corpo após strip do prefixo de produto.
    bullet(doc, "case principal. Construí apontamento de horas.",
           negrito_prefixo="Consol: ")
    # Bullet sem prefixo: verbo deve ser o primeiro token.
    bullet(doc, "Construí serviço de exportação de relatórios em Java.")

    h2(doc, "Formação")
    linha_data(doc, "Bacharelado: UFPA", "Jan 2017 - Dez 2021")

    return doc


# ---- Teste 1: documento válido passa ----

def test_documento_valido_nao_levanta():
    doc = _doc_valido_minimo()
    # Não deve levantar nenhuma exceção.
    validar(doc)


# ---- Teste 2: tabela presente falha ----

def test_tabela_presente_levanta_assertionerror_mencionando_tabela():
    doc = _doc_valido_minimo()
    # Adiciona uma tabela de layout (proibida pelo ATS).
    doc.add_table(rows=1, cols=2)
    doc.tables[0].rows[0].cells[0].text = "esquerda"
    doc.tables[0].rows[0].cells[1].text = "direita"

    with pytest.raises(AssertionError) as excinfo:
        validar(doc)

    assert "tabela" in str(excinfo.value).lower()


# ---- Teste 3: em-dash presente falha ----

def test_em_dash_presente_levanta_assertionerror_mencionando_em_dash():
    doc = _doc_valido_minimo()
    # Adiciona parágrafo com em-dash (U+2014).
    doc.add_paragraph("Texto com em-dash proibido \u2014 aqui.")

    with pytest.raises(AssertionError) as excinfo:
        validar(doc)

    # Mensagem deve mencionar "em-dash" (com hífen simples).
    assert "em-dash" in str(excinfo.value).lower()


def test_en_dash_presente_levanta_assertionerror():
    doc = _doc_valido_minimo()
    # Adiciona parágrafo com en-dash (U+2013).
    doc.add_paragraph("Texto com en-dash proibido \u2013 aqui.")

    with pytest.raises(AssertionError) as excinfo:
        validar(doc)

    assert "en-dash" in str(excinfo.value).lower()


# ---- Teste 4: seção obrigatória ausente falha ----

def test_secao_obrigatoria_ausente_levanta_assertionerror():
    """Constrói um Document sem a seção PERFIL e verifica que falha."""
    doc = Document()
    style(doc)
    nome(doc, "Candidato Nome")
    cargo(doc, "Engenheiro de Software Pleno")
    contato(doc, [_TELEFONE, _EMAIL])
    # Sem PERFIL. Pulamos direto para HABILIDADES, EXPERIÊNCIA, FORMAÇÃO.
    h2(doc, "Habilidades")
    bullet(doc, "Java", negrito_prefixo="Backend: ")
    h2(doc, "Experiência")
    bullet(doc, "Construí sistema de exemplo.")
    h2(doc, "Formação")
    linha_data(doc, "Bacharelado: UFPA", "Jan 2017 - Dez 2021")

    with pytest.raises(AssertionError) as excinfo:
        validar(doc)

    assert "PERFIL" in str(excinfo.value)


# ---- Teste 5: verbo de ação ausente em bullet de experiência falha ----

def test_verbo_de_acao_ausente_em_experiencia_levanta_mencionando_verbo():
    """Bullet de experiência começando com 'Trabalhei' (não aceito) deve falhar."""
    doc = Document()
    style(doc)
    nome(doc, "Candidato Nome")
    cargo(doc, "Engenheiro de Software Pleno")
    contato(doc, [_TELEFONE, _EMAIL])

    h2(doc, "Perfil")
    paragrafo(doc, "Perfil sintético.")
    h2(doc, "Habilidades")
    bullet(doc, "Java", negrito_prefixo="Backend: ")
    h2(doc, "Experiência")
    linha_data(doc, "iUsecase", "Jan 2023 - Atual")
    # Verbo "Trabalhei" não está na lista de aceitos.
    bullet(doc, "Trabalhei em manutenção de sistema legado.")
    h2(doc, "Formação")
    linha_data(doc, "Bacharelado: UFPA", "Jan 2017 - Dez 2021")

    with pytest.raises(AssertionError) as excinfo:
        validar(doc)

    assert "verbo" in str(excinfo.value).lower()


def test_verbo_de_acao_aceito_em_experiencia_nao_levanta():
    """Cada verbo da lista de aceitos deve passar quando usado."""
    verbos = [
        "Construí", "Modelei", "Implementei", "Mantive", "Desenvolvi",
        "Integrei", "Estruturei", "Participei", "Colaborei", "Refatorei",
    ]
    for verbo in verbos:
        doc = _doc_valido_minimo()
        # Substitui o bullet de experiência por um começando com `verbo`.
        # Acha o último bullet "Construí ..." e troca o texto via recriação
        # é complexo; em vez disso, criamos um novo doc controlado.
        doc = Document()
        style(doc)
        nome(doc, "Candidato Nome")
        cargo(doc, "Engenheiro de Software Pleno")
        contato(doc, [_TELEFONE, _EMAIL])
        h2(doc, "Perfil")
        paragrafo(doc, "Perfil sintético.")
        h2(doc, "Habilidades")
        bullet(doc, "Java", negrito_prefixo="Backend: ")
        h2(doc, "Experiência")
        linha_data(doc, "iUsecase", "Jan 2023 - Atual")
        bullet(doc, f"{verbo} um sistema de teste com Java.")
        h2(doc, "Formação")
        linha_data(doc, "Bacharelado: UFPA", "Jan 2017 - Dez 2021")

        validar(doc)  # Não deve levantar.


def test_verbo_de_acao_apos_prefixo_produto_eh_aceito():
    """Bullet com prefixo de produto em negrito (ex.: 'Consol: ') seguido de
    texto que tem o verbo entre os primeiros tokens deve passar."""
    doc = _doc_valido_minimo()
    # O doc válido já tem um bullet "Consol: ..." cujo corpo começa com
    # substantivo, mas a regra aceita verbo nos 3 primeiros tokens. Para
    # este teste, criamos um novo doc com prefixo seguido imediatamente
    # de verbo.
    doc = Document()
    style(doc)
    nome(doc, "Candidato Nome")
    cargo(doc, "Engenheiro de Software Pleno")
    contato(doc, [_TELEFONE, _EMAIL])
    h2(doc, "Perfil")
    paragrafo(doc, "Perfil sintético.")
    h2(doc, "Habilidades")
    bullet(doc, "Java", negrito_prefixo="Backend: ")
    h2(doc, "Experiência")
    linha_data(doc, "iUsecase", "Jan 2023 - Atual")
    bullet(doc, "Implementei exportação CSV em Spring Boot.",
           negrito_prefixo="ProdutoX: ")
    h2(doc, "Formação")
    linha_data(doc, "Bacharelado: UFPA", "Jan 2017 - Dez 2021")

    validar(doc)  # Não deve levantar.


# ---- Teste 6: consistência com manifesto - IA False omitida ----

def test_manifesto_ia_false_sem_secao_ia_passa():
    doc = _doc_valido_minimo()  # não tem "IA COMO EIXO"
    manifesto = {"ia": False}
    validar(doc, manifesto)  # Não deve levantar.


def test_manifesto_ia_false_com_secao_ia_falha():
    doc = _doc_valido_minimo()
    h2(doc, "IA como Eixo de Estudo e Aplicação")
    bullet(doc, "Estudo de RAG.", negrito_prefixo="RAG: ")

    manifesto = {"ia": False}
    with pytest.raises(AssertionError) as excinfo:
        validar(doc, manifesto)

    msg = str(excinfo.value).lower()
    assert "ia" in msg
    assert "manifesto" in msg or "eixo" in msg


def test_manifesto_ia_true_sem_secao_ia_falha():
    doc = _doc_valido_minimo()  # sem "IA COMO EIXO"
    manifesto = {"ia": True}
    with pytest.raises(AssertionError) as excinfo:
        validar(doc, manifesto)

    assert "ia" in str(excinfo.value).lower()


def test_manifesto_ia_true_com_secao_ia_passa():
    doc = _doc_valido_minimo()
    h2(doc, "IA como Eixo de Estudo e Aplicação")
    manifesto = {"ia": True}
    validar(doc, manifesto)  # Não deve levantar.


def test_manifesto_idiomas_false_sem_secao_passa():
    doc = _doc_valido_minimo()  # sem IDIOMAS
    manifesto = {"idiomas": False}
    validar(doc, manifesto)


def test_manifesto_idiomas_true_sem_secao_falha():
    doc = _doc_valido_minimo()
    manifesto = {"idiomas": True}
    with pytest.raises(AssertionError) as excinfo:
        validar(doc, manifesto)

    assert "idioma" in str(excinfo.value).lower()


# ---- Teste 7: consistência com manifesto - case pedido ausente ----

def test_manifesto_case_pedido_ausente_falha():
    """Manifesto pede o produto 'Consol' mas o Document não tem 'Consol'."""
    doc = _doc_valido_minimo()
    # Constrói um doc sem mencionar 'Consol'.
    doc2 = Document()
    style(doc2)
    nome(doc2, "Candidato Nome")
    cargo(doc2, "Engenheiro de Software Pleno")
    contato(doc2, [_TELEFONE, _EMAIL])
    h2(doc2, "Perfil")
    paragrafo(doc2, "Perfil sintético.")
    h2(doc2, "Habilidades")
    bullet(doc2, "Java", negrito_prefixo="Backend: ")
    h2(doc2, "Experiência")
    linha_data(doc2, "iUsecase", "Jan 2023 - Atual")
    # Bullet sem mencionar 'Consol'.
    bullet(doc2, "Construí outro produto qualquer.")
    h2(doc2, "Formação")
    linha_data(doc2, "Bacharelado: UFPA", "Jan 2017 - Dez 2021")

    manifesto = {
        "experiencias": [
            {"arquivo": "experiencias/iusecase.yml", "cases": ["Consol"]}
        ]
    }
    with pytest.raises(AssertionError) as excinfo:
        validar(doc2, manifesto)

    assert "Consol" in str(excinfo.value)


def test_manifesto_case_pedido_presente_passa():
    doc = _doc_valido_minimo()  # contém "Consol" no bullet de prefixo
    manifesto = {
        "experiencias": [
            {"arquivo": "experiencias/iusecase.yml", "cases": ["Consol"]}
        ]
    }
    validar(doc, manifesto)  # Não deve levantar.


def test_manifesto_habilidades_bucket_pedido_ausente_falha():
    """Manifesto pede bucket 'Cloud' mas o Document não tem 'Cloud'."""
    doc = _doc_valido_minimo()  # só tem "Backend (JVM)"
    manifesto = {"habilidades_buckets": ["Cloud"]}
    with pytest.raises(AssertionError) as excinfo:
        validar(doc, manifesto)

    assert "Cloud" in str(excinfo.value)


def test_manifesto_habilidades_bucket_pedido_presente_passa():
    doc = _doc_valido_minimo()  # tem "Backend (JVM)"
    manifesto = {"habilidades_buckets": ["Backend (JVM)"]}
    validar(doc, manifesto)  # Não deve levantar.


# ---- Regras de honestidade (defesa em profundidade no DOCX) ----
# Estas regras já vivem em data/validate.py sobre o YAML. Aqui garantem
# que o render não injetou texto de outra fonte que viole as regras.

def test_honestidade_jooq_no_bullet_consol_falha():
    """jOOQ não pode aparecer no bullet de Consol no Document final."""
    doc = _doc_valido_minimo()
    # Adiciona bullet Consol com jOOQ (violação da regra de honestidade).
    h2(doc, "Experiência")
    bullet(doc, "Desenvolvi algo com jOOQ no backend.",
           negrito_prefixo="Consol: ")
    with pytest.raises(AssertionError) as excinfo:
        validar(doc)
    assert "jOOQ" in str(excinfo.value)
    assert "Consol" in str(excinfo.value)


def test_honestidade_live2u_deve_declarar_sys3_falha():
    """Bullet de Live2U deve mencionar Sys3 (backend externo)."""
    doc = _doc_valido_minimo()
    h2(doc, "Experiência")
    # Descrição que omite qualquer menção a Sys3 ou "externo".
    bullet(doc, "Integrei serviço de RAG sobre exames no frontend Angular.",
           negrito_prefixo="Live2U: ")
    with pytest.raises(AssertionError) as excinfo:
        validar(doc)
    assert "Sys3" in str(excinfo.value) or "externo" in str(excinfo.value)
