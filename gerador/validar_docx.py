"""Validador ATS de Document python-docx.

Camada de assertions executada sobre o `Document` gerado por
`gerador.montar.montar`, antes de `doc.save()`. Garante que o render
não introduziu regressão nas regras ATS:
- sem tabelas de layout;
- sem em-dashes nem en-dashes;
- seções obrigatórias presentes;
- contato presente;
- verbos de ação no passado em bullets de experiência;
- consistência com o manifesto de seleção (toggles, cases, buckets).

Mensagens PT-BR. Levanta AssertionError ruidosamente na primeira
violação para diagnóstico rápido.

Uso programático:

    from gerador.montar import montar
    from gerador.validar_docx import validar

    doc = montar("manifesto.json")
    validar(doc, manifesto)
    doc.save("saida.docx")

Uso via CLI (valida DOCX já salvo, sem checar manifesto):

    python3 -m gerador.validar_docx saida.docx
"""

from __future__ import annotations

import sys
from typing import Optional

from docx import Document
from docx.document import Document as _DocumentType


# ---- Constantes das regras ----

# Em-dash (U+2014) e en-dash (U+2013) proibidos. Mantidos como literais
# porque são exatamente o que a regra detecta.
EM_DASH = "\u2014"
EN_DASH = "\u2013"

# Seções obrigatórias em maiúsculas, conforme renderizado por h2().
# PERFIL, HABILIDADES, EXPERIÊNCIA, FORMAÇÃO sempre presentes no
# currículo canônico. IDIOMAS e IA são condicionais (toggles do
# manifesto).
SECOES_OBRIGATORIAS = ("PERFIL", "HABILIDADES", "EXPERIÊNCIA", "FORMAÇÃO")

# Contato real do brief (regra 4 do contrato do gerador). A validação
# garante que o render não dropou telefone nem email por descuido.
TELEFONE_ESPERADO = "(00) 00000-0000"
EMAIL_ESPERADO = "candidato.exemplo@dominio.com"

# Verbos de ação aceitos no início (ou nos 3 primeiros tokens) do corpo
# de bullets de experiência. Espelha a lista do `gerar_curriculo_ats.py`
# e adiciona "Refatorei" pedido no brief da Task 4.
VERBOS_ACEITOS = {
    "Construí",
    "Modelei",
    "Implementei",
    "Mantive",
    "Desenvolvi",
    "Integrei",
    "Estruturei",
    "Participei",
    "Colaborei",
    "Refatorei",
}

# Prefixos de produto conhecidos na fonte YAML. Aparecem em bullets de
# experiência como negrito (ex.: "Consol: case principal. Construí...").
# A regra do verbo precisa descontar o prefixo antes de procurar o
# verbo. Lista derivada de data/experiencias/*.yml e pode crescer.
PREFIXOS_PRODUTO = {
    "Consol:",
    "Apontamento:",
    "Live2U:",
    "Agronegócio:",
    "Fintech:",
    "Logística:",
    "Gamificação corporativa:",
    "Automotivo/industrial:",
    "Compliance:",
    "Educação:",
    "Transversais:",
}


# ---- API pública ----

def validar(
    doc: _DocumentType,
    manifesto: Optional[dict] = None,
) -> None:
    """Valida o Document python-docx contra regras ATS.

    Levanta AssertionError com mensagem PT-BR se alguma regra falha.
    `manifesto` é opcional: se fornecido, valida consistência entre o
    manifesto e o Document gerado (ex.: se manifesto pediu ia=False, a
    seção IA não deve aparecer).

    Args:
        doc: Document python-docx gerado por montar().
        manifesto: dicionário do manifesto de seleção (opcional). Quando
            None, apenas as regras gerais ATS são checadas.

    Raises:
        AssertionError: na primeira regra ATS violada, com mensagem
            PT-BR descrevendo o problema.
    """
    texto = "\n".join(p.text for p in doc.paragraphs)

    # 1. Sem tabelas de layout (ATS parseia mal).
    _checar_sem_tabelas(doc)

    # 2. Sem em-dashes nem en-dashes em qualquer parágrafo.
    _checar_sem_dashes(texto)

    # 3. Seções obrigatórias presentes.
    _checar_secoes_obrigatorias(texto)

    # 4. Contato presente.
    _checar_contato(texto)

    # 5. Verbos de ação no passado em bullets de experiência.
    _checar_verbos_experiencia(doc)

    # 6. Consistência com manifesto (se fornecido).
    if manifesto is not None:
        _checar_manifesto(texto, manifesto)


# ---- Implementação das regras ----

def _checar_sem_tabelas(doc: _DocumentType) -> None:
    """Regra 1: ATS não parseia tabelas de layout. Deve haver zero."""
    n = len(doc.tables)
    assert n == 0, (
        f"ATS proíbe tabelas de layout. Encontradas: {n} tabela(s). "
        f"Use parágrafos com tab stops (linha_data) no lugar."
    )


def _checar_sem_dashes(texto: str) -> None:
    """Regra 2: sem em-dashes (U+2014) nem en-dashes (U+2013).

    Mantém os literais unicode porque são exatamente o que a regra
    detecta. O código fonte não usa esses caracteres em nenhum outro
    lugar.
    """
    assert EM_DASH not in texto, (
        "Em-dash (U+2014) encontrado no texto. Substitua por ponto, "
        "vírgula, dois-pontos ou hífen simples com espaços."
    )
    assert EN_DASH not in texto, (
        "En-dash (U+2013) encontrado no texto. Use hífen simples com "
        "espaços (ex.: 'Jan 2023 - Atual')."
    )


def _checar_secoes_obrigatorias(texto: str) -> None:
    """Regra 3: seções obrigatórias como headers maiúsculos no texto."""
    for secao in SECOES_OBRIGATORIAS:
        assert secao in texto, (
            f"Seção obrigatória ausente: {secao}. "
            f"Verifique a ordem canônica em gerador.montar."
        )


def _checar_contato(texto: str) -> None:
    """Regra 4: telefone e email reais presentes (não dropados pelo render)."""
    assert TELEFONE_ESPERADO in texto, (
        f"Telefone de contato ausente: esperado {TELEFONE_ESPERADO!r}."
    )
    assert EMAIL_ESPERADO in texto, (
        f"E-mail de contato ausente: esperado {EMAIL_ESPERADO!r}."
    )


def _checar_verbos_experiencia(doc: _DocumentType) -> None:
    """Regra 5: bullets de experiência começam com verbo aceito.

    Identifica a seção EXPERIÊNCIA pelo header maiúsculo, itera
    parágrafos até o próximo header, e para cada bullet (style
    "List Bullet") checa que o corpo tem verbo de ação aceito entre os
    3 primeiros tokens. Se houver prefixo de produto em negrito
    (ex.: "Consol: "), o prefixo é descontado antes da checagem.

    Caso especial do YAML real: "Consol: case principal. Construí..."
    tem um marcador "case principal." antes do verbo. Aceitamos verbo
    em qualquer dos 3 primeiros tokens do corpo, replicando a lógica do
    `gerar_curriculo_ats.py:validar()` linhas 264-286.
    """
    dentro_experiencia = False
    for p in doc.paragraphs:
        texto_p = p.text.strip()
        # Header de seção: tudo maiúsculo e sem ser bullet/lista.
        if _e_header_secao(p):
            dentro_experiencia = (texto_p.upper() == "EXPERIÊNCIA")
            continue
        if not dentro_experiencia:
            continue
        # Só checamos bullets (style "List Bullet").
        if p.style.name != "List Bullet":
            continue
        _checar_verbo_no_bullet(texto_p)


def _e_header_secao(p) -> bool:
    """Heurística: parágrafo é um header de seção (h2 em maiúsculas).

    Considera header se o texto for curto (<= 60 chars), não for bullet
    e estiver totalmente em maiúsculas (após strip). Isto cobre os h2()
    de render.py (PERFIL, HABILIDADES, etc.) sem depender de estilo
    nomeado, já que o render atual não define estilo próprio para h2.
    """
    texto = p.text.strip()
    if not texto or len(texto) > 60:
        return False
    if p.style.name == "List Bullet":
        return False
    # h2 sempre renderiza em maiúsculas (ver render.h2).
    return texto == texto.upper() and any(c.isalpha() for c in texto)


def _checar_verbo_no_bullet(texto_p: str) -> None:
    """Verifica verbo de ação em um bullet de experiência.

    Desconta prefixo de produto ("Consol: ", etc.) e aceita verbo nos
    3 primeiros tokens do corpo restante. Mensagem PT-BR menciona a
    palavra 'verbo' e o início encontrado para diagnóstico.
    """
    corpo = texto_p
    for prefixo in PREFIXOS_PRODUTO:
        if corpo.startswith(prefixo):
            corpo = corpo[len(prefixo):].strip()
            break
    if not corpo:
        return  # bullet vazio, nada a checar (não deve ocorrer)
    # Tolerância: marcador "case principal." pode aparecer antes do
    # verbo. Olhamos os 3 primeiros tokens.
    palavras = corpo.replace(".", " ").split()
    primeiros = palavras[:3] if len(palavras) >= 3 else palavras
    primeira = palavras[0].rstrip(",.;:") if palavras else ""
    tem_verbo = any(w.rstrip(",.;:") in VERBOS_ACEITOS for w in primeiros)
    assert tem_verbo, (
        f"Bullet de experiência deve começar com verbo de ação no passado "
        f"(entre os 3 primeiros tokens). Verbos aceitos: "
        f"{sorted(VERBOS_ACEITOS)}. Início encontrado: {primeira!r}. "
        f"Texto do bullet: {texto_p[:80]!r}."
    )


def _checar_manifesto(texto: str, manifesto: dict) -> None:
    """Regra 6: consistência entre manifesto e Document gerado."""
    _checar_toggle(texto, manifesto, "ia", "IA COMO EIXO")
    _checar_toggle(texto, manifesto, "idiomas", "IDIOMAS")
    _checar_cases(texto, manifesto)
    _checar_buckets(texto, manifesto)


def _checar_toggle(
    texto: str,
    manifesto: dict,
    chave: str,
    marcador: str,
) -> None:
    """Toggle do manifesto: True exige marcador, False proíbe marcador."""
    ligado = bool(manifesto.get(chave, False))
    presente = marcador in texto.upper()
    if ligado and not presente:
        raise AssertionError(
            f"Manifesto pediu {chave}=True mas a seção {marcador!r} "
            f"não aparece no Document. Verifique o render de {chave} "
            f"em gerador.montar."
        )
    if (not ligado) and presente:
        raise AssertionError(
            f"Manifesto pediu {chave}=False mas a seção {marcador!r} "
            f"aparece no Document. Verifique o toggle de {chave} "
            f"em gerador.montar."
        )


def _checar_cases(texto: str, manifesto: dict) -> None:
    """Cada case pedido no manifesto deve aparecer (por nome de produto)."""
    for exp in manifesto.get("experiencias", []) or []:
        for produto in exp.get("cases", []) or []:
            assert str(produto) in texto, (
                f"Manifesto pediu o case {produto!r} mas o produto não "
                f"aparece no Document. Verifique o render de experiência "
                f"em gerador.montar."
            )


def _checar_buckets(texto: str, manifesto: dict) -> None:
    """Cada bucket pedido no manifesto deve aparecer (por rótulo)."""
    for rotulo in manifesto.get("habilidades_buckets", []) or []:
        assert str(rotulo) in texto, (
            f"Manifesto pediu o bucket de habilidades {rotulo!r} mas o "
            f"rótulo não aparece no Document. Verifique o render de "
            f"habilidades em gerador.montar."
        )


# ---- CLI ----

def _cli() -> int:
    """Valida um DOCX já salvo. Retorna 0 (ok) ou 1 (falha)."""
    if len(sys.argv) < 2:
        print("[FAIL] uso: python3 -m gerador.validar_docx <arquivo.docx>")
        return 1
    caminho = sys.argv[1]
    try:
        doc = Document(caminho)
    except Exception as e:
        print(f"[FAIL] não foi possível abrir {caminho!r}: {e}")
        return 1
    try:
        validar(doc)
    except AssertionError as e:
        print(f"[FAIL] {e}")
        return 1
    n_paragrafos = len(doc.paragraphs)
    n_tabelas = len(doc.tables)
    n_caracteres = sum(len(p.text) for p in doc.paragraphs)
    print("[OK] validação DOCX passou.")
    print(f"  - Parágrafos: {n_paragrafos}")
    print(f"  - Tabelas: {n_tabelas} (deve ser 0)")
    print(f"  - Caracteres: {n_caracteres}")
    return 0


if __name__ == "__main__":
    sys.exit(_cli())
