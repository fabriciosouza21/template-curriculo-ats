"""Testes dos helpers ATS em gerador.render.

Cobrem:
- cada helper de conteúdo adiciona ao menos 1 parágrafo;
- style define A4 (210x297 mm) e margens 2cm;
- nome produz texto em maiúsculas;
- bullet com negrito_prefixo produz run bold + run normal;
- linha_data produz texto contendo esquerda e direita.
"""

import pytest
from docx import Document

from gerador.render import (
    style,
    nome,
    cargo,
    contato,
    h2,
    paragrafo,
    bullet,
    linha_data,
)


# Helpers de conteúdo (style é de configuração, testado à parte).
HELPERS_CONTEUDO = [
    ("nome", lambda d: nome(d, "Candidato Nome")),
    ("cargo", lambda d: cargo(d, "Engenheiro de Software Pleno")),
    ("contato", lambda d: contato(d, ["a@b.com", "(91) 99999-9999", "Cidade, UF"])),
    ("h2", lambda d: h2(d, "Experiência")),
    ("paragrafo", lambda d: paragrafo(d, "Texto de corpo do parágrafo.")),
    ("bullet", lambda d: bullet(d, "corpo do bullet")),
    ("bullet_prefixo", lambda d: bullet(d, "corpo do bullet", negrito_prefixo="Bucket: ")),
    ("linha_data", lambda d: linha_data(d, "Empresa X", "Jan 2020 - Dez 2021")),
]


@pytest.mark.parametrize("nome_helper,fn", HELPERS_CONTEUDO)
def test_helper_adiciona_ao_menos_um_paragrafo(nome_helper, fn):
    doc = Document()
    antes = len(doc.paragraphs)
    fn(doc)
    assert len(doc.paragraphs) >= antes + 1, (
        f"helper {nome_helper} deveria adicionar ao menos 1 parágrafo"
    )


def test_style_configura_a4_e_margens_2cm():
    doc = Document()
    style(doc)
    sec = doc.sections[0]

    # A4: 210 x 297 mm.
    assert abs(sec.page_width.mm - 210) < 0.01
    assert abs(sec.page_height.mm - 297) < 0.01

    # Margens 2cm nos quatro lados.
    assert abs(sec.top_margin.cm - 2.0) < 0.01
    assert abs(sec.bottom_margin.cm - 2.0) < 0.01
    assert abs(sec.left_margin.cm - 2.0) < 0.01
    assert abs(sec.right_margin.cm - 2.0) < 0.01


def test_style_configura_fonte_e_espacamento_do_normal():
    doc = Document()
    style(doc)
    normal = doc.styles["Normal"]

    assert normal.font.name == "Calibri"
    assert abs(normal.font.size.pt - 10.5) < 0.01
    assert normal.paragraph_format.line_spacing == 1.1


def test_nome_produz_maiusculas():
    doc = Document()
    nome(doc, "Candidato Nome")
    p = doc.paragraphs[-1]
    assert p.text == "CANDIDATO NOME"
    assert p.text == "Candidato Nome".upper()


def test_h2_produz_maiusculas_e_bold():
    doc = Document()
    h2(doc, "Experiência")
    p = doc.paragraphs[-1]
    assert p.text == "EXPERIÊNCIA"
    assert any(r.bold for r in p.runs)


def test_contato_junta_partes_com_separador_pipe():
    doc = Document()
    contato(doc, ["a@b.com", "(91) 99999-9999"])
    p = doc.paragraphs[-1]
    assert p.text == "a@b.com | (91) 99999-9999"


def test_bullet_com_prefixo_tem_run_bold_e_run_normal():
    doc = Document()
    bullet(doc, "corpo do bullet", negrito_prefixo="Prefixo: ")
    p = doc.paragraphs[-1]
    runs = p.runs

    assert len(runs) >= 2
    assert runs[0].text == "Prefixo: "
    assert runs[0].bold is True
    assert runs[1].text == "corpo do bullet"
    assert runs[1].bold is not True
    assert "Prefixo: " in p.text
    assert "corpo do bullet" in p.text


def test_bullet_sem_prefixo_tem_um_run_nao_bold():
    doc = Document()
    bullet(doc, "corpo do bullet")
    p = doc.paragraphs[-1]
    runs = p.runs

    assert len(runs) == 1
    assert runs[0].text == "corpo do bullet"
    assert runs[0].bold is not True


def test_linha_data_contem_esquerda_e_direita():
    doc = Document()
    linha_data(doc, "Empresa X", "Jan 2020 - Dez 2021")
    p = doc.paragraphs[-1]
    assert "Empresa X" in p.text
    assert "Jan 2020 - Dez 2021" in p.text


def test_linha_data_esquerda_em_bold():
    doc = Document()
    linha_data(doc, "Empresa X", "Data")
    p = doc.paragraphs[-1]
    # Primeiro run (esquerda) deve ser bold.
    assert any(r.text == "Empresa X" and r.bold for r in p.runs)
